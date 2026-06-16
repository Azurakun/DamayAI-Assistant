"""
Part 2: Adds BAB II (Tinjauan Pustaka), BAB III (Deskripsi Sistem),
BAB IV (Eksperimen), BAB V (Penutup), and Daftar Pustaka.
Run AFTER generate_final_doc_part1.py.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

FILE = os.path.join(os.path.dirname(__file__), "DamayAI_Dokumen_PPA_Revised_Formatted.docx")
doc = Document(FILE)

# ── Helpers (same as Part 1) ──
def add_heading_centered(text, bold=True, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    return p

def add_heading_left(text, bold=True, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_subheading_left(text, bold=True, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    return p

def add_body(text, indent=True, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_body_no_indent(text, bold=False, italic=False):
    return add_body(text, indent=False, bold=bold, italic=italic)

def add_bullet(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run(f"•  {text}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_numbered_item(number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    run = p.add_run(f"{number}.\t{text}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_image_placeholder(caption, akses_disini=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if akses_disini:
        run = p.add_run("[Gambar — Silakan masukkan gambar di sini]")
    else:
        run = p.add_run("[Gambar — Placeholder Screenshot]")
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = cap.add_run(caption)
    r2.bold = True
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(11)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
    return table


# ═══════════════════════════════════════════
# BAB II — TINJAUAN PUSTAKA, LANDASAN TEORI, DAN DESKRIPSI SISTEM
# ═══════════════════════════════════════════
add_heading_centered("BAB II\nTINJAUAN PUSTAKA, LANDASAN TEORI,\nDAN DESKRIPSI SISTEM")

# 2.1 Tinjauan Pustaka / Penelitian Terkait
add_heading_left("2.1  Tinjauan Pustaka")

add_body(
    "Tinjauan pustaka berikut menyajikan kajian terhadap penelitian-penelitian terdahulu "
    "yang relevan dengan pengembangan sistem chatbot asisten virtual berbasis Tiered RAG. "
    "Kajian ini mencakup metode-metode utama yang digunakan dalam penelitian tersebut "
    "serta hasil yang diperoleh, sebagai landasan perbandingan dan justifikasi terhadap "
    "pendekatan yang diambil dalam proyek akhir ini."
)

doc.add_paragraph()  # spacer
add_body_no_indent("Tabel 2.1  Tinjauan Pustaka", bold=True)
doc.add_paragraph()  # spacer

tinjauan_rows = [
    ["1", "Patrick Lewis (2020)", "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks",
     "RAG (Retrieval-Augmented Generation)",
     "Menggabungkan retrieval dan generatif model untuk menghasilkan jawaban berbasis dokumen yang lebih akurat dibanding model generatif murni."],
    ["2", "Jeff Johnson (2019)", "Billion-scale Similarity Search with GPUs",
     "FAISS (Facebook AI Similarity Search)",
     "Mengembangkan library FAISS untuk pencarian kemiripan vektor skala besar secara efisien menggunakan GPU."],
    ["3", "J. Devlin (2019)", "BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding",
     "BERT (Bidirectional Encoder Representations from Transformers)",
     "Menunjukkan bahwa model bahasa bidirectional dapat meningkatkan performa pada berbagai tugas NLP secara signifikan."],
    ["4", "Gemini Team, Google (2024)", "Gemini: A Family of Highly Capable Multimodal Models",
     "Multimodal Large Language Model",
     "Gemini menunjukkan kemampuan terdepan pada benchmark multimodal dan tugas pemrosesan bahasa alami."],
    ["5", "Shervin Minaee (2024)", "Large Language Models: A Survey",
     "Literature Survey",
     "Menyajikan survey komprehensif tentang perkembangan LLM, termasuk arsitektur, pelatihan, dan aplikasi."],
    ["6", "Woosuk Kwon (2023)", "Efficient Memory Management for Large Language Model Serving with PagedAttention",
     "PagedAttention / vLLM",
     "Mengusulkan PagedAttention untuk manajemen memori yang efisien pada serving LLM, meningkatkan throughput secara signifikan."],
    ["7", "Yunfan Gao (2023)", "Retrieval-Augmented Generation for Large Language Models: A Survey",
     "Literature Survey on RAG",
     "Survey komprehensif tentang teknik RAG, mencakup arsitektur, strategi retrieval, dan tantangan yang ada."],
    ["8", "The Vicuna Team (2023)", "Vicuna: An Open-Source Chatbot Impressing GPT-4 with 90% ChatGPT Quality",
     "Fine-tuning LLaMA pada data percakapan",
     "Vicuna mencapai kualitas 90% mendekati ChatGPT, menunjukkan potensi open-source LLM."],
]
add_table(["No", "Penulis / Tahun", "Judul Penelitian", "Metode", "Hasil"], tinjauan_rows)

add_body(
    "Berdasarkan tinjauan pustaka di atas, dapat disimpulkan bahwa pendekatan RAG telah "
    "banyak diteliti dan terbukti efektif dalam meningkatkan akurasi jawaban LLM dengan "
    "memanfaatkan sumber pengetahuan eksternal. Namun demikian, belum terdapat penelitian "
    "yang secara khusus mengimplementasikan pendekatan Tiered RAG dengan mekanisme prioritas "
    "data bertingkat (Memory Bank, Data Manual, Data Scraping) pada konteks sekolah menengah "
    "kejuruan di Indonesia. Proyek akhir ini mengisi celah penelitian tersebut dengan "
    "mengembangkan DamayAI Assistant yang mengadopsi arsitektur Tiered RAG untuk memastikan "
    "jawaban yang dihasilkan selalu mengutamakan sumber data yang paling terverifikasi."
)

add_body(
    "Selain itu, dari perspektif implementasi, sebagian besar penelitian terdahulu berfokus "
    "pada lingkungan universitas di luar negeri, sementara adaptasi teknologi ini untuk "
    "ekosistem SMK di Indonesia—dengan karakteristik data dan kebutuhan informasi yang "
    "berbeda—masih sangat terbatas. DamayAI Assistant dirancang untuk menjembatani "
    "kesenjangan tersebut dengan menyediakan solusi yang spesifik, kontekstual, dan "
    "berbasis data resmi SMKN 2 Indramayu."
)

# 2.2 Landasan Teori
add_heading_left("2.2  Landasan Teori")

# 2.2.1
add_subheading_left("2.2.1  Kecerdasan Buatan dan Pemrosesan Bahasa Alami")
add_body(
    "Kecerdasan buatan (Artificial Intelligence/AI) merupakan cabang ilmu komputer yang "
    "berfokus pada pengembangan sistem yang mampu melakukan tugas-tugas yang biasanya "
    "memerlukan kecerdasan manusia, seperti pembelajaran, penalaran, pengenalan pola, "
    "dan pengambilan keputusan. Dalam konteks pemrosesan bahasa alami (Natural Language "
    "Processing/NLP), kecerdasan buatan memungkinkan mesin untuk memahami, menafsirkan, "
    "dan menghasilkan bahasa manusia secara otomatis."
)
add_body(
    "Pemrosesan Bahasa Alami (NLP) adalah sub-bidang AI yang menggabungkan teknik "
    "linguistik komputasional, pembelajaran mesin, dan pembelajaran mendalam untuk "
    "memungkinkan komputer memproses dan memahami bahasa manusia dalam skala besar. "
    "Perkembangan NLP telah mengalami evolusi signifikan, mulai dari pendekatan berbasis "
    "aturan (rule-based) pada era awal, statistik pada tahun 1990-an, hingga pendekatan "
    "berbasis deep learning yang saat ini mendominasi."
)
add_body(
    "Dalam konteks pengembangan chatbot, NLP memainkan peran krusial dalam memungkinkan "
    "sistem memahami maksud pengguna (intent recognition), mengekstraksi entitas penting "
    "(entity extraction), dan menghasilkan respons yang koheren dan kontekstual. Teknologi "
    "NLP modern yang didukung oleh model transformer telah memungkinkan chatbot untuk tidak "
    "hanya memahami pertanyaan secara literal, tetapi juga menangkap nuansa kontekstual dan "
    "semantik dari percakapan."
)

# 2.2.2
add_subheading_left("2.2.2  Large Language Models (LLM)")
add_body(
    "Large Language Model (LLM) adalah model pembelajaran mendalam yang dilatih pada korpus "
    "teks berskala sangat besar menggunakan arsitektur transformer. LLM mampu memahami dan "
    "menghasilkan teks bahasa manusia dengan kualitas yang mendekati atau bahkan menyamai "
    "kemampuan manusia dalam berbagai tugas NLP, termasuk penerjemahan, peringkasan, "
    "penjawaban pertanyaan, dan pembuatan konten kreatif."
)
add_body(
    "Arsitektur transformer yang diperkenalkan oleh Vaswani et al. (2017) menjadi fondasi "
    "utama bagi pengembangan LLM. Mekanisme self-attention pada transformer memungkinkan "
    "model untuk menangkap hubungan kontekstual antar kata dalam kalimat tanpa terbatas "
    "oleh jarak antar kata, yang merupakan keterbatasan utama pada arsitektur sebelumnya "
    "seperti RNN dan LSTM. Model-model LLM terkemuka seperti GPT, BERT, LLaMA, dan Gemini "
    "semuanya dibangun berdasarkan varian arsitektur transformer ini."
)
add_body(
    "LLM bekerja dengan cara memprediksi token berikutnya dalam sebuah urutan teks "
    "berdasarkan konteks yang diberikan. Proses ini memungkinkan model untuk menghasilkan "
    "teks yang koheren dan relevan secara autoregresif. Namun, LLM memiliki keterbatasan "
    "inheren berupa pengetahuan yang terbatas pada data pelatihan (knowledge cutoff) dan "
    "kecenderungan untuk menghasilkan informasi yang tidak akurat atau fiktif, yang dikenal "
    "sebagai halusinasi (hallucination). Keterbatasan ini menjadi motivasi utama "
    "pengembangan pendekatan Retrieval-Augmented Generation (RAG)."
)

# 2.2.2.1
add_subheading_left("2.2.2.1  Google Gemini")
add_body(
    "Google Gemini adalah keluarga model multimodal yang dikembangkan oleh Google DeepMind, "
    "dirancang untuk memproses dan menghasilkan konten dalam berbagai modalitas, termasuk "
    "teks, gambar, audio, dan kode. Gemini tersedia dalam beberapa varian ukuran—Ultra, Pro, "
    "dan Flash—yang masing-masing dirancang untuk menyeimbangkan antara kemampuan dan "
    "efisiensi komputasi."
)
add_body(
    "Dalam proyek akhir ini, Google Gemini digunakan sebagai model generatif utama dalam "
    "arsitektur RAG. Gemini dipilih karena kemampuannya yang unggul dalam pemahaman dan "
    "generasi bahasa Indonesia, ketersediaan API yang mudah diakses melalui Google AI Studio, "
    "serta dukungan terhadap konteks percakapan multi-turn yang memungkinkan interaksi yang "
    "lebih natural. Model Gemini Pro yang digunakan dalam DamayAI memiliki kemampuan reasoning "
    "yang kuat dan mampu menghasilkan respons yang koheren berdasarkan konteks yang diberikan "
    "oleh hasil retrieval."
)

# 2.2.3
add_subheading_left("2.2.3  Retrieval-Augmented Generation (RAG)")
add_body(
    "Retrieval-Augmented Generation (RAG) adalah pendekatan yang menggabungkan mekanisme "
    "retrieval (pencarian informasi) dengan kemampuan generatif LLM untuk menghasilkan "
    "jawaban yang lebih akurat dan terverifikasi. Konsep RAG pertama kali diperkenalkan oleh "
    "Lewis et al. (2020) dan sejak itu telah menjadi salah satu paradigma paling berpengaruh "
    "dalam pengembangan sistem tanya-jawab berbasis AI."
)
add_body(
    "Dalam arsitektur RAG, ketika pengguna mengajukan pertanyaan, sistem terlebih dahulu "
    "melakukan pencarian semantik (semantic search) terhadap basis pengetahuan yang telah "
    "diindeks untuk menemukan dokumen atau potongan teks yang paling relevan. Hasil "
    "retrieval ini kemudian digabungkan dengan pertanyaan asli pengguna sebagai konteks "
    "tambahan (augmented context) yang diberikan kepada LLM. Dengan demikian, LLM tidak "
    "hanya mengandalkan pengetahuan parametriknya, tetapi juga mendapatkan akses terhadap "
    "informasi aktual dari sumber eksternal yang telah diverifikasi."
)
add_body(
    "Proses RAG secara umum terdiri dari tiga tahap utama, yaitu: (1) Indexing—pembuatan "
    "indeks vektor dari dokumen sumber melalui proses chunking, embedding, dan penyimpanan "
    "di vector database; (2) Retrieval—pencarian dokumen yang paling relevan berdasarkan "
    "kemiripan semantik antara query pengguna dan dokumen yang diindeks; dan (3) Generation—"
    "proses pembuatan jawaban oleh LLM berdasarkan konteks yang diperoleh dari hasil "
    "retrieval."
)

# 2.2.3.1
add_subheading_left("2.2.3.1  Keunggulan RAG dibanding Fine-tuning")
add_body(
    "Dalam konteks penyesuaian LLM untuk domain spesifik, terdapat dua pendekatan utama: "
    "fine-tuning dan RAG. Fine-tuning melibatkan pelatihan ulang model pada dataset "
    "domain-specific untuk mengubah pengetahuan parametrik model, sedangkan RAG "
    "mempertahankan model dasar dan menyediakan pengetahuan domain melalui mekanisme "
    "retrieval."
)
add_body("RAG memiliki beberapa keunggulan signifikan dibanding fine-tuning dalam konteks pengembangan chatbot informasi sekolah, antara lain sebagai berikut.")
add_bullet("RAG memungkinkan pembaruan pengetahuan secara real-time tanpa perlu melatih ulang model—cukup dengan menambahkan atau memperbarui dokumen di knowledge base.")
add_bullet("RAG menyediakan transparansi yang lebih baik karena jawaban dapat ditelusuri kembali ke sumber dokumen spesifik.")
add_bullet("RAG lebih efisien secara komputasional karena tidak memerlukan proses pelatihan yang mahal.")
add_bullet("RAG mengurangi risiko halusinasi karena LLM diarahkan untuk menghasilkan jawaban berdasarkan konteks yang disediakan, bukan mengandalkan pengetahuan parametrik yang mungkin sudah usang atau tidak akurat.")

# 2.2.4
add_subheading_left("2.2.4  Metode Pengumpulan Data untuk Sistem RAG")
add_body(
    "Dalam pengembangan sistem RAG, kualitas dan kelengkapan data yang digunakan sebagai "
    "knowledge base sangat menentukan kualitas jawaban yang dihasilkan. DamayAI Assistant "
    "menggunakan tiga metode pengumpulan data utama yang masing-masing memiliki karakteristik "
    "dan tingkat prioritas yang berbeda dalam mekanisme Tiered RAG."
)

# 2.2.4.1
add_subheading_left("2.2.4.1  Web Crawling dan Web Scraping")
add_body(
    "Web crawling adalah proses otomatis untuk menavigasi dan mengunduh halaman-halaman web "
    "secara sistematis, dimulai dari seed URL tertentu. Web scraping merupakan teknik "
    "lanjutan dari web crawling yang tidak hanya mengunduh halaman web, tetapi juga "
    "mengekstraksi informasi spesifik dari konten HTML. Dalam konteks DamayAI, web scraping "
    "digunakan untuk mengekstraksi informasi dari website resmi SMKN 2 Indramayu "
    "(smkn2indramayu.sch.id), termasuk data profil sekolah, program keahlian, data staf dan "
    "guru, serta informasi kegiatan sekolah."
)
add_body(
    "Proses web scraping pada DamayAI dilakukan menggunakan library BeautifulSoup dan "
    "Selenium untuk menangani halaman web yang bersifat statis maupun dinamis. Data yang "
    "diperoleh melalui scraping kemudian dibersihkan (cleaning), dinormalisasi, dan disimpan "
    "ke MongoDB sebelum akhirnya dipecah menjadi chunk dan di-embed ke FAISS. Data scraping "
    "memiliki prioritas terendah dalam mekanisme Tiered RAG karena kemungkinan adanya "
    "informasi yang belum terverifikasi atau tidak terbaru."
)

# 2.2.4.2
add_subheading_left("2.2.4.2  Pengunggahan Dokumen Manual")
add_body(
    "Pengunggahan dokumen manual (manual upload) merupakan metode pengumpulan data di mana "
    "administrator mengunggah dokumen resmi sekolah dalam format PDF, DOCX, atau PPTX ke "
    "dalam sistem. Metode ini menghasilkan data dengan tingkat kepercayaan yang lebih tinggi "
    "dibanding scraping karena dokumen yang diunggah umumnya merupakan dokumen resmi yang "
    "telah diverifikasi oleh pihak sekolah, seperti profil lengkap staf, pedoman sekolah, "
    "dan dokumen akreditasi."
)
add_body(
    "Proses pengunggahan dokumen manual pada DamayAI melibatkan ekstraksi teks dari berbagai "
    "format file menggunakan library seperti PyPDF2 untuk PDF, python-docx untuk DOCX, dan "
    "python-pptx untuk PPTX. Teks yang diekstraksi kemudian melalui proses chunking dan "
    "embedding yang sama dengan data dari sumber lain. Data manual memiliki prioritas "
    "menengah dalam mekanisme Tiered RAG—lebih tinggi dari scraping namun lebih rendah dari "
    "Memory Bank."
)

# 2.2.4.3
add_subheading_left("2.2.4.3  Memory Bank (Pembelajaran Berbasis Percakapan)")
add_body(
    "Memory Bank adalah fitur yang memungkinkan administrator untuk memasukkan pasangan "
    "pertanyaan-jawaban (Q&A pairs) secara langsung ke dalam sistem. Berbeda dengan "
    "scraping dan upload manual yang mengandalkan proses retrieval semantik, Memory Bank "
    "memungkinkan kurasi jawaban yang paling tepat dan terverifikasi untuk "
    "pertanyaan-pertanyaan spesifik yang sering diajukan (frequently asked questions)."
)
add_body(
    "Memory Bank memiliki prioritas tertinggi dalam mekanisme Tiered RAG karena setiap "
    "pasangan Q&A telah dikurasi secara manual oleh administrator yang memahami konteks dan "
    "kebutuhan informasi sekolah. Ketika pengguna mengajukan pertanyaan yang memiliki "
    "kemiripan semantik tinggi dengan entri di Memory Bank, sistem akan mengutamakan "
    "jawaban dari sumber ini terlebih dahulu sebelum mencari di lapisan data lainnya. "
    "Pendekatan ini memastikan bahwa jawaban yang paling tepercaya dan relevan selalu "
    "diprioritaskan."
)

# 2.2.5
add_subheading_left("2.2.5  Text Embedding dan Vector Database")
add_subheading_left("2.2.5.1  Text Embedding")
add_body(
    "Text embedding adalah proses mengkonversi teks menjadi representasi vektor numerik "
    "berdimensi tinggi yang menangkap makna semantik dari teks tersebut. Model embedding "
    "memetakan teks ke dalam ruang vektor sedemikian rupa sehingga teks dengan makna yang "
    "serupa berada berdekatan dalam ruang vektor tersebut. Teknik ini menjadi fondasi dari "
    "pencarian semantik (semantic search) yang digunakan dalam sistem RAG."
)
add_body(
    "Dalam DamayAI, digunakan model embedding dari Google (models/embedding-001) yang "
    "menghasilkan vektor berdimensi 768 untuk setiap input teks. Model ini dipilih karena "
    "kompatibilitasnya dengan ekosistem Google AI, kualitas embedding yang baik untuk bahasa "
    "Indonesia, dan efisiensi komputasional yang memadai untuk skala data sekolah. Proses "
    "embedding dilakukan pada setiap chunk teks yang dihasilkan dari proses chunking, dan "
    "vektor-vektor hasil embedding disimpan dalam FAISS untuk pencarian kemiripan yang cepat."
)

add_subheading_left("2.2.5.2  FAISS (Facebook AI Similarity Search)")
add_body(
    "FAISS (Facebook AI Similarity Search) adalah library open-source yang dikembangkan oleh "
    "Facebook AI Research untuk pencarian kemiripan (similarity search) dan pengelompokan "
    "(clustering) vektor berdimensi tinggi secara efisien. FAISS mendukung berbagai metode "
    "indexing yang memungkinkan pencarian miliaran vektor dengan latensi yang sangat rendah."
)
add_body(
    "Dalam arsitektur DamayAI, FAISS berfungsi sebagai vector store yang menyimpan seluruh "
    "embedding dari dokumen yang telah diproses. Ketika pengguna mengajukan pertanyaan, "
    "query tersebut di-embed menggunakan model yang sama, dan FAISS melakukan pencarian "
    "nearest neighbor untuk menemukan chunk-chunk yang paling mirip secara semantik. DamayAI "
    "menggunakan IndexFlatL2 yang melakukan pencarian exhaustif berbasis jarak Euclidean "
    "(L2), yang cocok untuk skala data sekolah yang tidak terlalu besar namun memerlukan "
    "akurasi pencarian yang tinggi."
)

# 2.2.6
add_subheading_left("2.2.6  Text Chunking (Pemecahan Teks)")
add_body(
    "Text chunking adalah proses memecah dokumen panjang menjadi potongan-potongan teks "
    "yang lebih kecil (chunk) sebelum dilakukan embedding. Proses ini diperlukan karena "
    "model embedding memiliki batasan panjang input, dan chunk yang lebih kecil menghasilkan "
    "pencarian semantik yang lebih akurat karena setiap chunk mewakili topik atau konsep "
    "yang lebih fokus."
)
add_body(
    "DamayAI menggunakan strategi fixed-size chunking dengan overlap, di mana teks dipecah "
    "menjadi chunk berukuran sekitar 500 karakter dengan overlap 100 karakter antar chunk. "
    "Overlap ini memastikan bahwa konteks yang melintasi batas chunk tidak hilang. Selain "
    "itu, dilakukan juga pembersihan teks (text cleaning) yang mencakup penghapusan karakter "
    "tidak perlu, normalisasi whitespace, dan penanganan karakter khusus sebelum proses "
    "chunking dilakukan."
)

# 2.2.7
add_subheading_left("2.2.7  MongoDB sebagai Basis Data Dokumen")
add_body(
    "MongoDB adalah sistem manajemen basis data NoSQL berorientasi dokumen yang menyimpan "
    "data dalam format BSON (Binary JSON). MongoDB dipilih dalam pengembangan DamayAI karena "
    "kemampuannya dalam menyimpan dan mengelola data semi-terstruktur dan tidak terstruktur "
    "yang fleksibel, sesuai dengan karakteristik data sekolah yang beragam format dan "
    "skemanya."
)
add_body(
    "Dalam arsitektur DamayAI, MongoDB berfungsi sebagai basis data utama yang menyimpan "
    "metadata dokumen, log percakapan, data pengguna administrator, dan data Memory Bank. "
    "MongoDB juga menyimpan data mentah hasil scraping dan upload dokumen sebelum diproses "
    "lebih lanjut. Keunggulan MongoDB dalam hal fleksibilitas skema dan kemampuan horizontal "
    "scaling menjadikannya pilihan yang tepat untuk sistem yang memerlukan penyimpanan data "
    "dengan format yang bervariasi."
)

# 2.2.8
add_subheading_left("2.2.8  Flask sebagai Framework Backend")
add_body(
    "Flask adalah micro-framework Python untuk pengembangan aplikasi web yang bersifat ringan "
    "dan modular. Flask dipilih sebagai framework backend DamayAI karena kesederhanaannya, "
    "fleksibilitas dalam pengembangan API, serta dukungan ekosistem Python yang luas untuk "
    "library-library AI dan NLP."
)
add_body(
    "Dalam arsitektur DamayAI, Flask menangani seluruh logika backend, termasuk routing API "
    "untuk chatbot, autentikasi administrator, pengelolaan knowledge base (scraping dan "
    "upload manual), pengelolaan Memory Bank, dan penyajian halaman web. Flask juga berperan "
    "sebagai penghubung antara komponen-komponen sistem seperti MongoDB, FAISS, dan Google "
    "Gemini API. Penggunaan Flask memungkinkan arsitektur yang bersih dan mudah di-maintain, "
    "dengan pemisahan yang jelas antara logika bisnis, data access, dan presentation layer."
)

# 2.2.9
add_subheading_left("2.2.9  Aspek Keamanan dalam Sistem Berbasis LLM")
add_body(
    "Pengembangan sistem yang melibatkan LLM dan data institusional memerlukan perhatian "
    "khusus terhadap aspek keamanan. Dalam konteks DamayAI, terdapat beberapa aspek "
    "keamanan yang diperhatikan, yaitu sebagai berikut."
)
add_bullet("Keamanan akses—hanya administrator yang terautentikasi yang dapat mengakses panel pengelolaan knowledge base dan Memory Bank.")
add_bullet("Validasi input—mencegah injection attack melalui sanitasi input pengguna.")
add_bullet("Proteksi data—memastikan data internal sekolah tidak bocor melalui respons chatbot yang tidak semestinya.")
add_bullet("Rate limiting—mencegah penyalahgunaan sistem melalui pembatasan jumlah request.")
add_body(
    "Selain itu, aspek keamanan juga mencakup pencegahan prompt injection, di mana pengguna "
    "berupaya memanipulasi prompt sistem untuk mengakses informasi yang tidak diizinkan atau "
    "mengubah perilaku chatbot. DamayAI menerapkan sanitasi input dan system prompt yang "
    "ketat untuk meminimalkan risiko serangan ini. Kombinasi mekanisme keamanan ini "
    "memastikan bahwa sistem dapat dioperasikan secara aman dalam lingkungan institusi "
    "pendidikan."
)

# 2.3 Kerangka Berpikir
add_heading_left("2.3  Kerangka Berpikir")
add_body(
    "Kerangka berpikir dalam penelitian ini dibangun berdasarkan identifikasi permasalahan "
    "aksesibilitas informasi di SMKN 2 Indramayu yang kemudian diformulasikan menjadi solusi "
    "berbasis teknologi kecerdasan buatan. Alur berpikir dimulai dari identifikasi kondisi "
    "eksisting layanan informasi sekolah yang masih konvensional, kemudian dilakukan analisis "
    "kebutuhan pengguna (siswa, guru, staf TU), hingga perancangan solusi berupa chatbot "
    "asisten virtual yang mengimplementasikan pendekatan Tiered RAG. Secara sistematis, "
    "kerangka berpikir dapat diuraikan dalam tahapan-tahapan berikut."
)
add_numbered_item(1, "Identifikasi masalah: informasi tersebar, respons manual lambat, AI publik tidak memiliki data internal sekolah.")
add_numbered_item(2, "Analisis kebutuhan: sistem terpusat, real-time, berbasis data terverifikasi, dapat diakses kapan saja.")
add_numbered_item(3, "Perancangan solusi: chatbot berbasis web dengan arsitektur Tiered RAG menggunakan tiga lapisan knowledge base (Memory Bank, Data Manual, Data Scraping).")
add_numbered_item(4, "Implementasi: pengembangan sistem menggunakan Flask, MongoDB, FAISS, dan Google Gemini AI.")
add_numbered_item(5, "Pengujian: evaluasi fungsionalitas, akurasi jawaban, dan performa sistem melalui black box testing dan pengujian respons.")

doc.add_paragraph()
add_image_placeholder("Gambar 2.6  Kerangka Berpikir Penelitian", akses_disini=True)

# 2.4 Deskripsi Permasalahan
add_heading_left("2.4  Deskripsi Permasalahan")
add_body(
    "Berdasarkan identifikasi permasalahan yang telah diuraikan pada Bab I, deskripsi "
    "permasalahan dalam proyek akhir ini dapat dirinci sebagai berikut. SMKN 2 Indramayu "
    "sebagai institusi pendidikan yang besar dengan delapan program keahlian menghadapi "
    "tantangan signifikan dalam pengelolaan dan penyajian informasi kepada pemangku "
    "kepentingan, terutama siswa baru dan calon peserta didik."
)
add_body(
    "Permasalahan pertama berkaitan dengan fragmentasi sumber informasi. Saat ini, informasi "
    "sekolah tersebar di berbagai kanal yang tidak terintegrasi, mulai dari mading fisik yang "
    "hanya dapat diakses di lingkungan sekolah, grup WhatsApp yang tidak terorganisir dan "
    "sulit dicari, hingga website resmi sekolah yang tidak selalu diperbarui secara berkala. "
    "Kondisi ini menyebabkan siswa baru, yang merupakan pengguna informasi terbesar, "
    "kesulitan menemukan jawaban atas pertanyaan mereka secara mandiri dan efisien."
)
add_body(
    "Permasalahan kedua berkaitan dengan ketidaksesuaian solusi AI publik. Meskipun asisten "
    "AI seperti ChatGPT telah tersedia secara luas, sistem-sistem ini tidak memiliki akses "
    "terhadap data internal SMKN 2 Indramayu. Akibatnya, jawaban yang diberikan seringkali "
    "tidak relevan, tidak akurat, atau bahkan menyesatkan ketika digunakan untuk menanyakan "
    "informasi spesifik tentang sekolah, seperti nama guru, jadwal kegiatan, atau prosedur "
    "PPDB."
)
add_body(
    "Permasalahan ketiga berkaitan dengan beban kerja staf Tata Usaha. Setiap hari, staf TU "
    "harus menjawab puluhan pertanyaan serupa dari siswa baru dan orang tua, mulai dari "
    "pertanyaan tentang jam operasional sekolah, syarat PPDB, hingga lokasi ruang guru. "
    "Proses manual ini tidak hanya membebani staf TU tetapi juga menyebabkan keterlambatan "
    "respons dan ketidakkonsistenan informasi yang diberikan."
)

# 2.5 Deskripsi Solusi
add_heading_left("2.5  Deskripsi Solusi")
add_body(
    "Untuk mengatasi permasalahan yang telah diidentifikasi, proyek akhir ini mengusulkan "
    "solusi berupa pengembangan DamayAI Assistant, sebuah chatbot asisten virtual berbasis "
    "web yang mengimplementasikan pendekatan Tiered Retrieval-Augmented Generation. Solusi "
    "ini dirancang dengan mempertimbangkan karakteristik unik kebutuhan informasi di "
    "lingkungan SMK dan memanfaatkan teknologi AI terkini untuk memberikan layanan informasi "
    "yang akurat, cepat, dan terverifikasi."
)
add_body(
    "Solusi utama yang ditawarkan adalah implementasi mekanisme Tiered RAG yang membagi "
    "knowledge base menjadi tiga lapisan dengan prioritas yang berbeda. Lapisan pertama dan "
    "tertinggi adalah Memory Bank, yang berisi pasangan Q&A yang telah dikurasi manual oleh "
    "administrator dan memiliki kepercayaan tertinggi. Lapisan kedua adalah Data Manual, "
    "yang berisi dokumen resmi sekolah yang diunggah oleh administrator dalam format PDF, "
    "DOCX, atau PPTX. Lapisan ketiga dan terendah adalah Data Scraping, yang diperoleh "
    "melalui proses web scraping dari website resmi sekolah."
)
add_body(
    "Mekanisme prioritas ini bekerja dengan cara: ketika pengguna mengajukan pertanyaan, "
    "sistem terlebih dahulu mencari jawaban di Memory Bank. Jika ditemukan kecocokan yang "
    "memadai, jawaban dari Memory Bank langsung digunakan. Jika tidak, sistem melanjutkan "
    "pencarian ke Data Manual, dan terakhir ke Data Scraping. Pendekatan ini memastikan "
    "bahwa jawaban yang paling terverifikasi selalu diprioritaskan, sekaligus meminimalkan "
    "risiko halusinasi AI yang dapat terjadi ketika LLM menghasilkan jawaban tanpa dasar "
    "data yang kuat."
)
add_body(
    "Secara teknis, DamayAI dibangun menggunakan Flask sebagai framework backend, MongoDB "
    "sebagai basis data dokumen, FAISS sebagai vector store untuk pencarian semantik, dan "
    "Google Gemini sebagai model generatif. Sistem menyediakan dua antarmuka utama: "
    "antarmuka chatbot untuk pengguna umum (siswa, guru, masyarakat) dan antarmuka panel "
    "admin untuk pengelolaan knowledge base dan Memory Bank."
)

# 2.6 Desain Sistem
add_heading_left("2.6  Desain Sistem")

add_subheading_left("2.6.1  Arsitektur Sistem")
add_body(
    "Arsitektur sistem DamayAI Assistant mengadopsi pola arsitektur three-tier yang terdiri "
    "dari presentation layer, business logic layer, dan data access layer. Presentation "
    "layer berupa antarmuka web yang dapat diakses melalui browser, business logic layer "
    "ditangani oleh Flask backend, dan data access layer melibatkan MongoDB dan FAISS "
    "sebagai komponen penyimpanan data."
)
add_body(
    "Pada arsitektur ini, ketika pengguna mengirimkan pertanyaan melalui antarmuka chatbot, "
    "request dikirim ke Flask backend yang kemudian melakukan proses retrieval: pertanyaan "
    "di-embed menggunakan Google embedding model, pencarian dilakukan di FAISS vector store "
    "berdasarkan mekanisme tiered priority, konteks yang diperoleh digabungkan dengan "
    "pertanyaan asli dan dikirim ke Google Gemini API untuk generasi jawaban. Jawaban yang "
    "dihasilkan kemudian dikembalikan ke pengguna melalui antarmuka web."
)

doc.add_paragraph()
add_image_placeholder("Gambar 2.1  Desain Arsitektur Sistem (Akses Disini)", akses_disini=True)

add_body(
    "Alur pengumpulan dan pemrosesan data DamayAI melibatkan tiga sumber data utama yang "
    "masing-masing melalui proses yang berbeda sebelum masuk ke knowledge base. Web "
    "crawling/scraping mengekstraksi data dari website resmi sekolah, upload manual "
    "memproses dokumen yang diunggah administrator, dan Memory Bank menyimpan pasangan Q&A "
    "yang dikurasi langsung. Seluruh data kemudian disimpan di MongoDB, dipecah menjadi "
    "chunk, di-embed, dan diindeks di FAISS untuk keperluan retrieval."
)

doc.add_paragraph()
add_image_placeholder("Gambar 2.2  Alur Pengumpulan dan Pemrosesan Data DamayAI (Akses Disini)", akses_disini=True)

add_subheading_left("2.6.2  Pemodelan Fungsi Sistem")
add_body(
    "Pemodelan fungsi sistem DamayAI menggunakan diagram use case untuk menggambarkan "
    "interaksi antara aktor dan fungsi-fungsi sistem. Terdapat dua aktor utama: Pengguna "
    "Umum (siswa, guru, masyarakat) dan Administrator. Pengguna Umum dapat berinteraksi "
    "dengan chatbot untuk bertanya, melihat riwayat chat, dan memberi feedback. "
    "Administrator memiliki akses tambahan untuk login/logout, mengelola knowledge base "
    "(scraping dan upload manual), mengelola Memory Bank, melakukan uji coba AI, dan "
    "melihat laporan bug."
)

doc.add_paragraph()
add_image_placeholder("Gambar 2.3  Use Case Diagram (Akses Disini)", akses_disini=True)

add_subheading_left("2.6.3  Pemodelan Data dan Proses")
add_body(
    "Pemodelan data dan proses menggunakan Data Flow Diagram (DFD) untuk menggambarkan "
    "aliran data dalam sistem. DFD Level 0 menunjukkan sistem DamayAI sebagai satu proses "
    "tunggal dengan entitas eksternal berupa Pengguna Umum, Staf TU/Guru, dan Administrator. "
    "Data store utama adalah Knowledge Base yang terdiri dari MongoDB dan FAISS."
)

doc.add_paragraph()
add_image_placeholder("Gambar 2.4  DFD Level 0 (Akses Disini)", akses_disini=True)

add_body(
    "DFD Level 1 memecah proses utama menjadi empat sub-proses: (1) Autentikasi—menangani "
    "login dan verifikasi administrator; (2) Knowledge Base Management—menangani "
    "pengelolaan data melalui scraping, upload manual, dan Memory Bank; (3) Inferensi & "
    "Retrieval—menangani proses pencarian dan generasi jawaban; dan (4) Pelaporan "
    "Bug—menangani pencatatan dan pelaporan bug yang ditemukan pengguna."
)

doc.add_paragraph()
add_image_placeholder("Gambar 2.5  DFD Level 1 (Akses Disini)", akses_disini=True)

# Save checkpoint
doc.save(FILE)
print("Part 2 complete: BAB II added.")
