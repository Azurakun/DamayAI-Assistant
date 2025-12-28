import os
import json
import shutil
import google.generativeai as genai
from flask import Flask, request, jsonify, Response, stream_with_context, send_from_directory
from dotenv import load_dotenv
from scraper import scrape_from_file, extract_text_from_pdf, extract_text_from_pptx
from vector_store import create_vector_db, get_retrievers
from database import (
    init_db, add_scraped_data, get_all_scraped_data, update_scraped_data, delete_scraped_data,
    add_to_memory, get_all_memory_data, update_memory_data, delete_memory_data,
    add_manual_data, get_all_manual_data, update_manual_data, delete_manual_data,
    add_bug_report, get_all_bug_reports, update_bug_report_status, delete_bug_report,
    get_db # Ensure this is imported if used directly
)
from werkzeug.utils import secure_filename
import docx
from io import BytesIO
import datetime

# Load env variables
load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

UPLOADS_DIR = os.path.join(os.path.dirname(__file__), '..', 'uploads')
os.makedirs(os.path.join(UPLOADS_DIR, 'bugs'), exist_ok=True)
os.makedirs("db", exist_ok=True)

# Initialize MongoDB
try:
    init_db()
except Exception as e:
    print(f"Failed to initialize database: {e}")

app = Flask(__name__, static_folder='../frontend', static_url_path='/')
model = genai.GenerativeModel(model_name="gemini-2.5-flash")

FAISS_MEMORY_PATH = "db/faiss_index_memory"
FAISS_MANUAL_PATH = "db/faiss_index_manual"
FAISS_SCRAPED_PATH = "db/faiss_index_scraped"

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'pptx'}
ALLOWED_BUG_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'mp4', 'mov', 'avi'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def allowed_bug_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_BUG_EXTENSIONS

def extract_text_from_docx(file_stream):
    """Extracts text and tables from a DOCX file stream, converting tables to Markdown."""
    try:
        stream_buffer = BytesIO(file_stream.read())
        document = docx.Document(stream_buffer)
        
        content_parts = []
        for element in document.element.body:
            if element.tag.endswith('p'):
                content_parts.append(docx.text.paragraph.Paragraph(element, document).text)
            elif element.tag.endswith('tbl'):
                table = docx.table.Table(element, document)
                header = [cell.text.replace('\n', ' ').strip() for cell in table.rows[0].cells]
                md_table = f"\n| {' | '.join(header)} |\n"
                md_table += f"| {' | '.join(['---'] * len(header))} |\n"
                for row in table.rows[1:]:
                    row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                    md_table += f"| {' | '.join(row_data)} |\n"
                content_parts.append(md_table)
        
        return "\n".join(content_parts)
    except Exception:
        file_stream.seek(0) 
        document = docx.Document(file_stream)
        return "\n".join([p.text for p in document.paragraphs])

@app.route('/api/report_bug', methods=['POST'])
def report_bug_handler():
    try:
        description = request.form.get('description')
        file = request.files.get('file')
        if not description:
            return jsonify({"status": "error", "message": "Deskripsi masalah harus diisi."}), 400
        file_path = None
        if file and allowed_bug_file(file.filename):
            filename = secure_filename(file.filename)
            relative_path = os.path.join('bugs', filename)
            save_path = os.path.join(UPLOADS_DIR, 'bugs', filename)
            file.save(save_path)
            file_path = relative_path.replace(os.path.sep, '/')
        add_bug_report(description, file_path)
        return jsonify({"status": "success", "message": "Laporan bug berhasil dikirim."})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/api/get_bug_reports', methods=['GET'])
def get_bug_reports_handler():
    reports = get_all_bug_reports()
    return jsonify(reports)

@app.route('/api/bug_reports/<string:report_id>/status', methods=['PUT'])
def update_bug_status_handler(report_id):
    try:
        data = request.json
        new_status = data.get('status')
        if not new_status:
            return jsonify({"status": "error", "message": "Status is required."}), 400
        update_bug_report_status(report_id, new_status)
        return jsonify({"status": "success", "message": f"Bug report {report_id} status updated."})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/api/bug_reports/<string:report_id>', methods=['DELETE'])
def delete_bug_handler(report_id):
    try:
        delete_bug_report(report_id)
        return jsonify({"status": "success", "message": f"Bug report {report_id} deleted."})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/api/add_manual_text', methods=['POST'])
def add_manual_text_handler():
    try:
        data = request.json
        title = data.get('title', 'Tanpa Judul')
        content = data.get('content', '')
        if not content.strip():
            return jsonify({"status": "error", "message": "Konten teks tidak boleh kosong."}), 400
        
        source_name = f"manual-text-{title}"
        add_manual_data(source_name=source_name, title=title, content=content)
        return jsonify({"status": "success", "message": "Data teks manual berhasil ditambahkan."})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/api/add_manual_file', methods=['POST'])
def add_manual_file_handler():
    try:
        title = request.form.get('title', '')
        file = request.files.get('file')

        if not file or not allowed_file(file.filename):
            return jsonify({"status": "error", "message": "File tidak valid atau tidak disediakan."}), 400

        filename = secure_filename(file.filename)
        final_title = title if title.strip() else os.path.splitext(filename)[0]
        
        content = ""
        ext = filename.rsplit('.', 1)[1].lower()
        if ext == 'pdf':
            content = extract_text_from_pdf(file.stream)
        elif ext == 'docx':
            content = extract_text_from_docx(file.stream)
        elif ext == 'pptx':
            content = extract_text_from_pptx(file.stream)
        elif ext == 'txt':
            content = file.stream.read().decode('utf-8')
        
        if not content or not content.strip():
            return jsonify({"status": "error", "message": "Gagal mengekstrak teks atau file kosong."}), 500

        source_name = f"manual-file-{filename}"
        add_manual_data(source_name=source_name, title=final_title, content=content)
        return jsonify({"status": "success", "message": f"Konten dari file '{filename}' berhasil ditambahkan."})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/api/save_memory', methods=['POST'])
def save_memory_handler():
    try:
        data = request.json
        question = data.get('question')
        answer = data.get('answer')
        if not question or not answer:
            return jsonify({"status": "error", "message": "Pertanyaan dan jawaban harus diisi."}), 400
        add_to_memory(question, answer)
        return jsonify({"status": "success", "message": "Percakapan berhasil disimpan ke memori."})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/api/chat', methods=['POST'])
def chat_handler():
    data = request.json
    user_query = data.get('query', '')
    history = data.get('history', [])
    final_answer = "Maaf, terjadi kesalahan saat memproses permintaan Anda."
    for thought in generate_response(user_query, history):
        try:
            if thought.get('step') == 'final_answer':
                final_answer = thought.get('data', final_answer)
        except AttributeError:
            continue
    return jsonify({"response": final_answer})

@app.route('/api/admin_chat', methods=['POST'])
def admin_chat_handler():
    data = request.json
    user_query = data.get('query', '')
    history = data.get('history', [])
    return Response(stream_with_context(generate_response_stream(user_query, history)), mimetype='application/x-ndjson')

def generate_response_stream(user_query, history):
    for thought in generate_response(user_query, history):
        yield json.dumps(thought) + '\n'

def generate_response(user_query, history):
    if not user_query:
        yield {"step": "error", "data": "Query tidak boleh kosong."}
        return
        
    yield {"step": "start", "data": f"Menerima pertanyaan: '{user_query}'"}

    retriever_memory, retriever_manual, retriever_scraped = get_retrievers()
    
    # Store explicit list of retrieved info for prompt construction
    retrieved_knowledge = []

    # TAHAP 1: Cari di Memory Bank
    yield {"step": "memory_search", "data": "Mencari di Memory Bank..."}
    if retriever_memory:
        try:
            docs = retriever_memory.invoke(user_query)
            if docs:
                yield {"step": "memory_found", "data": f"{len(docs)} dokumen relevan ditemukan di Memory Bank."}
                for doc in docs:
                    retrieved_knowledge.append({
                        "source_type": "Memory Bank",
                        "title": doc.metadata.get('title', 'Unknown'),
                        "source": "Memory Bank", # Generic source for memory
                        "content": doc.page_content
                    })
            else:
                yield {"step": "memory_not_found", "data": "Tidak ada yang cocok di Memory Bank."}
        except Exception as e:
            yield {"step": "error", "data": f"Error saat mencari di Memory Bank: {e}"}

    # TAHAP 2: Cari di Data Manual 
    yield {"step": "manual_search", "data": "Mencari di Data Manual..."}
    if retriever_manual:
        try:
            docs = retriever_manual.invoke(user_query)
            if docs:
                yield {"step": "manual_found", "data": f"{len(docs)} dokumen relevan ditemukan di Data Manual."}
                for doc in docs:
                     retrieved_knowledge.append({
                        "source_type": "Data Manual",
                        "title": doc.metadata.get('title', 'Unknown'),
                        "source": doc.metadata.get('source', 'Manual Upload'),
                        "content": doc.page_content
                    })
            else:
                yield {"step": "manual_not_found", "data": "Tidak ada yang cocok di Data Manual."}
        except Exception as e:
            yield {"step": "error", "data": f"Error saat mencari di Data Manual: {e}"}

    # TAHAP 3: Cari di Data Scraping
    yield {"step": "scrape_search", "data": "Mencari di Data Scraping..."}
    if retriever_scraped:
        try:
            docs = retriever_scraped.invoke(user_query)
            if docs:
                yield {"step": "scrape_found", "data": f"{len(docs)} dokumen relevan ditemukan di Data Scraping."}
                for doc in docs:
                     retrieved_knowledge.append({
                        "source_type": "Website Scraping",
                        "title": doc.metadata.get('title', 'Website'),
                        "source": doc.metadata.get('source', '#'), # Should be URL
                        "content": doc.page_content
                    })
            else:
                yield {"step": "scrape_not_found", "data": "Tidak ada yang cocok di Data Scraping."}
        except Exception as e:
            yield {"step": "error", "data": f"Error saat mencari di Data Scraping: {e}"}
    
    # Send retrieving info to frontend for debug (optional)
    if retrieved_knowledge:
        debug_info = [{"source": f"[{item['source_type']}] {item['title']}", "content": item['content'][:100]+"..."} for item in retrieved_knowledge]
        yield {"step": "retrieved_docs", "data": debug_info}

    # Construct the Context String with strict Source labeling
    context_str = ""
    for i, item in enumerate(retrieved_knowledge):
        context_str += f"""
        --- DOCUMENT #{i+1} ---
        [Tipe]: {item['source_type']}
        [Judul]: {item['title']}
        [Source/URL]: {item['source']}
        [Konten]:
        {item['content']}
        -----------------------
        """

    yield {"step": "final_prompt", "data": "Menyusun jawaban akhir..."}
    try:
        final_prompt_text = f"""
        ### SYSTEM PROMPT (DamayAI) ###

        # Identitas & Gaya Bicara
        Anda adalah DamayAI, asisten digital SMKN 2 Indramayu.
        - **Human-like**: Bicaralah secara luwes, natural, dan sopan seperti manusia (Adik Panca/Dik Panca).
        - **Fleksibel**: Anda BOLEH mengobrol santai (small talk) tanpa data database jika pengguna hanya menyapa atau bertanya kabar.
        - **Grounding Wajib**: JIKA pengguna bertanya tentang fakta, info sekolah, atau data teknis, Anda WAJIB menggunakan "DATA PENDUKUNG" di bawah.
        - **Jujur**: Jika data tidak ada di context, katakan belum tahu, tapi tetaplah ramah.

        # Aturan Sitasi (PENTING!)
        Agar pengguna bisa melihat sumber data, ikuti aturan ini saat mengambil fakta dari "DATA PENDUKUNG":
        1. Ambil informasi dari dokumen.
        2. Di akhir kalimat/paragraf yang relevan, tambahkan tag sitasi khusus ini:
           `[CITE: Source/URL | Judul Dokumen]`
        3. Contoh: "Pendaftaran dibuka bulan Mei [CITE: https://smkn2-im.sch.id/daftar | Info PPDB]."
        4. JANGAN membuat link Markdown sendiri `[Judul](URL)`, gunakan format `[CITE:...]` saja. Frontend yang akan mengubahnya menjadi tombol (chip).

        # Format Jawaban
        1. Gunakan **Markdown** (Bold `**`, Italic `*`, List `-`, Tabel `|...|`).
        2. Buat jawaban ringkas, padat, dan mudah dibaca (poin-poin sangat disarankan).
        3. Sertakan `[IMAGE: url]` jika ada gambar di data pendukung.

        ---
        # DATA PENDUKUNG (Gunakan ini untuk fakta)
        {context_str if context_str else "Tidak ada data spesifik ditemukan. Gunakan pengetahuan umum hanya untuk sapaan/obrolan ringan."}
        
        # PERCAKAPAN
        Riwayat: {history}
        User Query: "{user_query}"
        ---
        
        Jawaban (Ingat tag [CITE:...] jika menggunakan data):
        """
        
        chat_session = model.start_chat(history=history)
        final_response = chat_session.send_message(final_prompt_text)
        
        yield {"step": "final_answer", "data": final_response.text}
    except Exception as e:
        yield {"step": "error", "data": f"Gagal menghasilkan jawaban akhir. Error: {e}"}


@app.route('/api/delete_faiss', methods=['POST'])
def delete_faiss_handler():
    try:
        paths = [FAISS_MEMORY_PATH, FAISS_MANUAL_PATH, FAISS_SCRAPED_PATH]
        deleted_count = 0
        for path in paths:
            if os.path.exists(path):
                shutil.rmtree(path)
                deleted_count += 1
        
        if deleted_count > 0:
            os.makedirs("db", exist_ok=True)
            return jsonify({"status": "success", "message": f"Berhasil menghapus {deleted_count} direktori indeks FAISS."})
        else:
            return jsonify({"status": "info", "message": "Tidak ada direktori indeks FAISS yang ditemukan untuk dihapus."})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/api/delete_db', methods=['POST'])
def delete_db_handler():
    try:
        db_to_drop = get_db()
        db_to_drop.scraped_data.drop()
        db_to_drop.manual_data.drop()
        db_to_drop.memory_bank.drop()
        init_db()
        return jsonify({"status": "success", "message": "Semua koleksi database berhasil dikosongkan."})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/api/scrape', methods=['POST'])
def scrape_handler():
    def generate_logs():
        urls_file = 'urls_to_scrape.txt'
        yield f"Membaca file '{urls_file}'...\n"
        for result in scrape_from_file(urls_file):
            if result['status'] == 'success':
                add_scraped_data(result['url'], result['title'], result['content'], result.get('image_url'))
                yield f"BERHASIL: {result['url']} - {result['title']}\n"
            else:
                yield f"DILEWATI/ERROR: {result['url']} - {result['reason']}\n"
    return Response(stream_with_context(generate_logs()), mimetype='text/plain')

@app.route('/api/reindex', methods=['POST'])
def reindex_handler():
    return Response(stream_with_context(create_vector_db()), mimetype='text/plain')

@app.route('/api/get-data', methods=['GET'])
def get_data_handler():
    scraped = get_all_scraped_data()
    manual = get_all_manual_data()
    memory = get_all_memory_data()

    all_data = []
    for item in scraped:
        item['type'] = 'Scrap'
        item['timestamp'] = item.get('scraped_at')
        all_data.append(item)
        
    for item in manual:
        item['type'] = 'Manual'
        item['timestamp'] = item.get('added_at')
        item['url'] = item.get('source_name')
        all_data.append(item)

    for item in memory:
        item['type'] = 'Memory'
        item['timestamp'] = item.get('saved_at')
        item['title'] = item.get('question')
        item['content'] = item.get('answer')
        item['url'] = f"Memory Bank #{item.get('id')}"
        all_data.append(item)

    all_data_sorted = sorted(all_data, key=lambda x: x.get('timestamp') or datetime.datetime.min, reverse=True)
    return jsonify(all_data_sorted)

@app.route('/api/data/<string:type>/<string:item_id>', methods=['PUT', 'DELETE'])
def update_delete_data_handler(type, item_id):
    try:
        if request.method == 'PUT':
            data = request.json
            if type == 'Scrap':
                update_scraped_data(item_id, data.get('title'), data.get('content'))
            elif type == 'Manual':
                update_manual_data(item_id, data.get('title'), data.get('content'))
            elif type == 'Memory':
                update_memory_data(item_id, data.get('title'), data.get('content'))
            else:
                return jsonify({"status": "error", "message": "Tipe data tidak valid."}), 400
            return jsonify({"status": "success", "message": "Data berhasil diperbarui."})

        elif request.method == 'DELETE':
            if type == 'Scrap':
                delete_scraped_data(item_id)
            elif type == 'Manual':
                delete_manual_data(item_id)
            elif type == 'Memory':
                delete_memory_data(item_id)
            else:
                return jsonify({"status": "error", "message": "Tipe data tidak valid."}), 400
            return jsonify({"status": "success", "message": "Data berhasil dihapus."})
            
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/')
def serve_index():
    return send_from_directory('../frontend', 'index.html')

@app.route('/admin')
def serve_admin():
    return send_from_directory('../frontend', 'admin.html')

@app.route('/uploads/<path:filename>')
def serve_upload(filename):
    return send_from_directory(UPLOADS_DIR, filename)
    
@app.route('/<path:path>')
def serve_static(path):
    return send_from_directory('../frontend', path)

if __name__ == '__main__':
    app.run(debug=True, port=5000)