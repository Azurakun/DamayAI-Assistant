"""
Script to generate the revised DamayAI PPA document (Bab 1-3)
following proper academic report structure.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

# ── Default Style ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# ── Helper Functions ──
def add_heading_centered(text, level=1, bold=True, size=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    return p

def add_heading_left(text, level=2, bold=True, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body(text, indent=True, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_body_no_indent(text, bold=False, italic=False):
    return add_body(text, indent=False, bold=bold, italic=italic)

def add_bullet(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run(f"•  {text}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_numbered_item(number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    run = p.add_run(f"{number}.\t{text}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_blank_lines(n=1):
    for _ in range(n):
        doc.add_paragraph()

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
    return table

def add_image_placeholder(caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Gambar — Placeholder Screenshot]")
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = cap.add_run(caption)
    r2.bold = True
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(11)


# ═══════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════
add_blank_lines(2)
add_heading_centered("PROPOSAL", size=14)
add_blank_lines(1)
add_heading_centered("LAPORAN AKHIR", size=14)
add_blank_lines(1)
add_heading_centered(
    "IMPLEMENTASI CHATBOT ASISTEN VIRTUAL SMKN 2 INDRAMAYU\n"
    "MENGGUNAKAN METODE TIERED RETRIEVAL-AUGMENTED GENERATION\n"
    "DENGAN MEKANISME PRIORITAS DATA", size=14)
add_blank_lines(2)
add_heading_centered("Achmad Arditio Sumartono", size=12)
add_heading_centered("NRP. 3124510063", size=12)
add_blank_lines(2)
add_heading_centered("DOSEN PEMBIMBING", size=12)
add_blank_lines(1)
add_heading_centered("Prof. Dr. Arna Fariza, S.Kom., M.Kom.", size=12)
add_heading_centered("NIP. 197107081999032001", size=11)
add_blank_lines(1)
add_heading_centered("Fitrah Maharani Humaira, M.Kom.", size=12)
add_heading_centered("NIP.", size=11)
add_blank_lines(3)
add_heading_centered(
    "PROGRAM STUDI DIPLOMA TIGA\n"
    "TEKNIK INFORMATIKA\n"
    "DEPARTEMEN TEKNIK INFORMATIKA DAN KOMPUTER\n"
    "POLITEKNIK ELEKTRONIKA NEGERI SURABAYA\n"
    "2025", size=12)

doc.add_page_break()

# ═══════════════════════════════════════════
# DAFTAR ISI (placeholder)
# ═══════════════════════════════════════════
add_heading_centered("DAFTAR ISI", size=14)
add_body_no_indent("(Daftar Isi akan di-generate secara otomatis oleh Microsoft Word)")
doc.add_page_break()

add_heading_centered("DAFTAR GAMBAR", size=14)
add_body_no_indent("(Daftar Gambar akan di-generate secara otomatis oleh Microsoft Word)")
doc.add_page_break()

add_heading_centered("DAFTAR TABEL", size=14)
add_body_no_indent("(Daftar Tabel akan di-generate secara otomatis oleh Microsoft Word)")
doc.add_page_break()


# ═══════════════════════════════════════════
# BAB I — PENDAHULUAN
# ═══════════════════════════════════════════
add_heading_centered("BAB I\nPENDAHULUAN", size=14)

# 1.1 Latar Belakang
add_heading_left("1.1  Latar Belakang")

add_body(
    "SMK Negeri 2 Indramayu merupakan salah satu sekolah kejuruan negeri unggulan di "
    "Kabupaten Indramayu, Jawa Barat, yang menyelenggarakan berbagai program keahlian "
    "berorientasi industri, mulai dari NKPI, APHPI, TJKT, PPLG, Teknik Alat Berat, "
    "Kuliner, APAPL, hingga TPFL. Dengan jumlah peserta didik yang terus bertumbuh dari "
    "tahun ke tahun, sekolah ini memiliki ekosistem informasi yang sangat luas, mulai dari "
    "jadwal kegiatan, data staf dan guru, informasi PPDB, hingga rekrutmen kerja yang "
    "difasilitasi oleh Bursa Kerja Khusus (BKK). Tingginya volume informasi ini tentu "
    "menjadi tantangan tersendiri bagi pengelolaan layanan informasi sekolah yang selama "
    "ini masih dilakukan secara konvensional dan tidak terpusat."
)

add_body(
    "Perkembangan kecerdasan buatan, khususnya hadirnya Large Language Model (LLM) "
    "seperti Google Gemini, telah membuka peluang besar untuk mengotomasi layanan "
    "informasi semacam ini secara lebih efisien dan kontekstual. Meski begitu, pengelolaan "
    "informasi di SMKN 2 Indramayu saat ini masih tersebar di berbagai kanal yang tidak "
    "terpusat, seperti mading fisik, grup WhatsApp, hingga website yang tidak selalu "
    "diperbarui. Kondisi ini mempersulit siswa baru untuk menemukan informasi secara "
    "cepat dan mandiri, serta membebani staf Tata Usaha yang harus menjawab "
    "pertanyaan-pertanyaan serupa secara berulang setiap harinya."
)

add_body(
    "Penggunaan asisten AI publik seperti ChatGPT pun tidak dapat menjadi solusi yang "
    "tepat karena sistem tersebut tidak memiliki akses terhadap data internal sekolah, "
    "sehingga rentan memberikan informasi yang tidak akurat atau bahkan menyesatkan. "
    "Pendekatan Retrieval-Augmented Generation (RAG) telah terbukti mampu mengatasi "
    "permasalahan ini dengan menggabungkan kemampuan generatif LLM dan pencarian semantik "
    "berbasis vector store seperti FAISS, sehingga jawaban yang dihasilkan dapat didasarkan "
    "pada dokumen terverifikasi milik institusi."
)

add_body(
    "Berdasarkan kajian terhadap penelitian terdahulu, sistem chatbot berbasis RAG untuk "
    "institusi pendidikan telah banyak dikembangkan di lingkungan universitas luar negeri, "
    "namun belum ada yang secara spesifik menyasar konteks sekolah menengah kejuruan (SMK) "
    "di Indonesia, khususnya dengan pendekatan tiered retrieval tiga lapis yang membedakan "
    "antara Memory Bank, Data Manual, dan Data Scraping. Oleh karena itu, pada proyek "
    "akhir ini dikembangkanlah DamayAI Assistant, sebuah chatbot asisten virtual yang "
    "dibangun menggunakan pendekatan Tiered Retrieval-Augmented Generation berbasis data "
    "resmi SMKN 2 Indramayu, mencakup profil sekolah, data staf, kegiatan, dan informasi "
    "akademik, sehingga setiap jawaban dapat diprioritaskan berdasarkan sumber data yang "
    "paling tepercaya."
)

# 1.2 Identifikasi Permasalahan
add_heading_left("1.2  Identifikasi Permasalahan")

add_body(
    "Berdasarkan latar belakang yang telah diuraikan, terdapat tiga permasalahan utama "
    "yang diidentifikasi sebagai berikut."
)

add_numbered_item(1,
    "Informasi seputar SMKN 2 Indramayu, seperti data staf, jadwal kegiatan, profil "
    "sekolah, dan informasi PPDB, masih tersebar di berbagai kanal yang tidak terpusat, "
    "termasuk mading fisik, grup WhatsApp, dan website yang tidak selalu diperbarui, "
    "sehingga siswa baru kesulitan mengakses informasi secara cepat dan mandiri."
)
add_numbered_item(2,
    "Penggunaan asisten AI publik seperti ChatGPT tidak efektif sebagai solusi karena "
    "sistem tersebut tidak memiliki akses terhadap data internal SMKN 2 Indramayu, "
    "sehingga rentan menghasilkan jawaban yang tidak relevan, tidak akurat, atau bahkan "
    "keliru terkait informasi spesifik sekolah."
)
add_numbered_item(3,
    "Pertanyaan umum dari siswa baru saat ini dijawab secara manual oleh staf Tata Usaha "
    "setiap harinya, yang menyebabkan keterlambatan respons serta tidak adanya skalabilitas "
    "dalam pengelolaan informasi sekolah."
)

# 1.3 Tujuan
add_heading_left("1.3  Tujuan")

add_body(
    "Kegiatan proyek akhir ini membangun suatu sistem chatbot asisten virtual berbasis web "
    "untuk mengatasi permasalahan aksesibilitas informasi di lingkungan SMKN 2 Indramayu "
    "dengan memanfaatkan teknologi kecerdasan buatan. Adapun tujuan yang ingin dicapai "
    "secara spesifik adalah sebagai berikut."
)

add_numbered_item(1,
    "Menyediakan chatbot asisten virtual bernama DamayAI yang dapat menjawab pertanyaan "
    "seputar SMKN 2 Indramayu, mulai dari profil sekolah, data staf dan guru, hingga "
    "kegiatan sekolah, secara interaktif dan real-time tanpa perlu menghubungi staf "
    "secara manual."
)
add_numbered_item(2,
    "Membangun knowledge base kustom milik sekolah melalui proses scraping website dan "
    "penginputan data manual, sehingga setiap jawaban yang diberikan selalu berlandaskan "
    "informasi resmi dan spesifik SMKN 2 Indramayu."
)
add_numbered_item(3,
    "Mengimplementasikan mekanisme Tiered Retrieval-Augmented Generation yang "
    "memprioritaskan data terverifikasi milik sekolah sebelum menarik informasi dari "
    "sumber lain, guna meminimalkan halusinasi AI dan meningkatkan kepercayaan pengguna "
    "terhadap jawaban yang diberikan."
)

# 1.4 Manfaat
add_heading_left("1.4  Manfaat")

add_body(
    "Manfaat yang diharapkan dari pengembangan sistem DamayAI Assistant adalah sebagai "
    "berikut."
)

add_numbered_item(1,
    "Bagi siswa baru SMKN 2 Indramayu, DamayAI memberikan kemudahan dalam mengakses "
    "informasi seputar sekolah secara mandiri, cepat, dan kapan saja tanpa harus menunggu "
    "respons manual dari staf. Hal ini diharapkan dapat mempersingkat masa adaptasi siswa "
    "baru terhadap lingkungan sekolah."
)
add_numbered_item(2,
    "Bagi staf Tata Usaha dan guru SMKN 2 Indramayu, sistem ini membantu mengurangi beban "
    "kerja berulang akibat menjawab pertanyaan yang serupa setiap harinya, sehingga staf "
    "dapat lebih berfokus pada tugas-tugas yang membutuhkan penanganan langsung dan bersifat "
    "lebih kompleks."
)
add_numbered_item(3,
    "Bagi institusi SMKN 2 Indramayu secara keseluruhan, DamayAI dapat menjadi sarana "
    "digitalisasi layanan informasi sekolah yang modern, terpusat, dan berbasis data "
    "terverifikasi, sekaligus menjadi bukti penerapan teknologi kecerdasan buatan di "
    "lingkungan sekolah menengah kejuruan di Indonesia."
)
add_numbered_item(4,
    "Bagi pengembang dan komunitas akademik, proyek ini memberikan kontribusi berupa "
    "implementasi nyata pendekatan Tiered RAG untuk chatbot informasi sekolah menengah "
    "kejuruan di Indonesia, yang dapat dijadikan referensi atau dasar pengembangan sistem "
    "serupa di institusi pendidikan lainnya."
)

doc.add_page_break()

# ═══════════════════════════════════════════
# BAB II — TINJAUAN PUSTAKA, LANDASAN TEORI, DAN DESKRIPSI SISTEM
# ═══════════════════════════════════════════
add_heading_centered("BAB II\nTINJAUAN PUSTAKA, LANDASAN TEORI,\nDAN DESKRIPSI SISTEM", size=14)

# 2.1 Tinjauan Pustaka / Penelitian Terkait
add_heading_left("2.1  Tinjauan Pustaka")

add_body(
    "Tinjauan pustaka berikut menyajikan kajian terhadap penelitian-penelitian terdahulu "
    "yang relevan dengan pengembangan sistem chatbot asisten virtual berbasis Tiered RAG. "
    "Kajian ini mencakup metode-metode utama yang digunakan dalam penelitian tersebut "
    "serta hasil yang diperoleh, sebagai landasan perbandingan dan justifikasi terhadap "
    "pendekatan yang diambil dalam proyek akhir ini."
)

doc.add_paragraph()  # spacer
add_body_no_indent("Tabel 2.1  Tinjauan Pustaka", bold=True)
doc.add_paragraph()  # spacer

tinjauan_rows = [
    ["1", "Patrick Lewis (2020)", "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks",
     "RAG (Retrieval-Augmented Generation)",
     "Menggabungkan retrieval dan generatif model untuk menghasilkan jawaban berbasis dokumen yang lebih akurat dibanding model generatif murni."],
    ["2", "Jeff Johnson (2019)", "Billion-scale Similarity Search with GPUs",
     "FAISS (Facebook AI Similarity Search)",
     "Mengembangkan library FAISS untuk pencarian kemiripan vektor skala besar secara efisien menggunakan GPU."],
    ["3", "J. Devlin (2019)", "BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding",
     "BERT (Bidirectional Encoder Representations from Transformers)",
     "Menunjukkan bahwa model bahasa bidirectional dapat meningkatkan performa pada berbagai tugas NLP secara signifikan."],
    ["4", "Gemini Team, Google (2024)", "Gemini: A Family of Highly Capable Multimodal Models",
     "Multimodal Large Language Model",
     "Gemini menunjukkan kemampuan terdepan pada benchmark multimodal dan tugas pemrosesan bahasa alami."],
    ["5", "Shervin Minaee (2024)", "Large Language Models: A Survey",
     "Literature Survey",
     "Menyajikan survey komprehensif tentang perkembangan LLM, termasuk arsitektur, pelatihan, dan aplikasi."],
    ["6", "Woosuk Kwon (2023)", "Efficient Memory Management for Large Language Model Serving with PagedAttention",
     "PagedAttention / vLLM",
     "Mengusulkan PagedAttention untuk manajemen memori yang efisien pada serving LLM, meningkatkan throughput secara signifikan."],
    ["7", "Yunfan Gao (2023)", "Retrieval-Augmented Generation for Large Language Models: A Survey",
     "Literature Survey on RAG",
     "Survey komprehensif tentang teknik RAG, mencakup arsitektur, strategi retrieval, dan tantangan yang ada."],
    ["8", "The Vicuna Team (2023)", "Vicuna: An Open-Source Chatbot Impressing GPT-4 with 90% ChatGPT Quality",
     "Fine-tuning LLaMA pada data percakapan",
     "Vicuna mencapai kualitas 90% mendekati ChatGPT, menunjukkan potensi open-source LLM."],
]
add_table(["No", "Penulis / Tahun", "Judul Penelitian", "Metode", "Hasil"], tinjauan_rows)

add_body(
    "Berdasarkan tinjauan pustaka di atas, dapat disimpulkan bahwa pendekatan RAG telah "
    "banyak diteliti dan terbukti efektif dalam meningkatkan akurasi jawaban LLM dengan "
    "memanfaatkan sumber pengetahuan eksternal. Namun demikian, belum terdapat penelitian "
    "yang secara khusus mengimplementasikan pendekatan Tiered RAG dengan mekanisme prioritas "
    "data bertingkat (Memory Bank, Data Manual, Data Scraping) pada konteks sekolah menengah "
    "kejuruan di Indonesia. Proyek akhir ini mengisi celah penelitian tersebut dengan "
    "mengembangkan DamayAI Assistant yang mengadopsi arsitektur Tiered RAG untuk memastikan "
    "jawaban yang dihasilkan selalu mengutamakan sumber data yang paling terverifikasi."
)

add_body(
    "Selain itu, dari perspektif implementasi, sebagian besar penelitian terdahulu berfokus "
    "pada lingkungan universitas di luar negeri, sementara adaptasi teknologi ini untuk "
    "ekosistem SMK di Indonesia—dengan karakteristik data dan kebutuhan informasi yang "
    "berbeda—masih sangat terbatas. DamayAI Assistant dirancang untuk menjembatani "
    "kesenjangan tersebut dengan menyediakan solusi yang spesifik, kontekstual, dan "
    "berbasis data resmi SMKN 2 Indramayu."
)

# 2.2 Landasan Teori
add_heading_left("2.2  Landasan Teori")

# 2.2.1
add_heading_left("2.2.1  Kecerdasan Buatan dan Pemrosesan Bahasa Alami")
add_body(
    "Kecerdasan buatan (Artificial Intelligence/AI) merupakan cabang ilmu komputer yang "
    "berfokus pada pengembangan sistem yang mampu melakukan tugas-tugas yang biasanya "
    "memerlukan kecerdasan manusia, seperti pembelajaran, penalaran, pengenalan pola, "
    "dan pengambilan keputusan. Dalam konteks pemrosesan bahasa alami (Natural Language "
    "Processing/NLP), kecerdasan buatan memungkinkan mesin untuk memahami, menafsirkan, "
    "dan menghasilkan bahasa manusia secara otomatis."
)
add_body(
    "Pemrosesan Bahasa Alami (NLP) adalah sub-bidang AI yang menggabungkan teknik "
    "linguistik komputasional, pembelajaran mesin, dan pembelajaran mendalam untuk "
    "memungkinkan komputer memproses dan memahami bahasa manusia dalam skala besar. "
    "Perkembangan NLP telah mengalami evolusi signifikan, mulai dari pendekatan berbasis "
    "aturan (rule-based) pada era awal, statistik pada tahun 1990-an, hingga pendekatan "
    "berbasis deep learning yang saat ini mendominasi."
)
add_body(
    "Dalam konteks pengembangan chatbot, NLP memainkan peran krusial dalam memungkinkan "
    "sistem memahami maksud pengguna (intent recognition), mengekstraksi entitas penting "
    "(entity extraction), dan menghasilkan respons yang koheren dan kontekstual. Teknologi "
    "NLP modern yang didukung oleh model transformer telah memungkinkan chatbot untuk tidak "
    "hanya memahami pertanyaan secara literal, tetapi juga menangkap nuansa kontekstual dan "
    "semantik dari percakapan."
)

# 2.2.2
add_heading_left("2.2.2  Large Language Models (LLM)")
add_body(
    "Large Language Model (LLM) adalah model pembelajaran mendalam yang dilatih pada korpus "
    "teks berskala sangat besar menggunakan arsitektur transformer. LLM mampu memahami dan "
    "menghasilkan teks bahasa manusia dengan kualitas yang mendekati atau bahkan menyamai "
    "kemampuan manusia dalam berbagai tugas NLP, termasuk penerjemahan, peringkasan, "
    "penjawaban pertanyaan, dan pembuatan konten kreatif."
)
add_body(
    "Arsitektur transformer yang diperkenalkan oleh Vaswani et al. (2017) menjadi fondasi "
    "utama bagi pengembangan LLM. Mekanisme self-attention pada transformer memungkinkan "
    "model untuk menangkap hubungan kontekstual antar kata dalam kalimat tanpa terbatas "
    "oleh jarak antar kata, yang merupakan keterbatasan utama pada arsitektur sebelumnya "
    "seperti RNN dan LSTM. Model-model LLM terkemuka seperti GPT, BERT, LLaMA, dan Gemini "
    "semuanya dibangun berdasarkan varian arsitektur transformer ini."
)
add_body(
    "LLM bekerja dengan cara memprediksi token berikutnya dalam sebuah urutan teks "
    "berdasarkan konteks yang diberikan. Proses ini memungkinkan model untuk menghasilkan "
    "teks yang koheren dan relevan secara autoregresif. Namun, LLM memiliki keterbatasan "
    "inheren berupa pengetahuan yang terbatas pada data pelatihan (knowledge cutoff) dan "
    "kecenderungan untuk menghasilkan informasi yang tidak akurat atau fiktif, yang dikenal "
    "sebagai halusinasi (hallucination). Keterbatasan ini menjadi motivasi utama "
    "pengembangan pendekatan Retrieval-Augmented Generation (RAG)."
)

# 2.2.2.1
add_heading_left("2.2.2.1  Google Gemini")
add_body(
    "Google Gemini adalah keluarga model multimodal yang dikembangkan oleh Google DeepMind, "
    "dirancang untuk memproses dan menghasilkan konten dalam berbagai modalitas, termasuk "
    "teks, gambar, audio, dan kode. Gemini tersedia dalam beberapa varian ukuran—Ultra, Pro, "
    "dan Flash—yang masing-masing dirancang untuk menyeimbangkan antara kemampuan dan "
    "efisiensi komputasi."
)
add_body(
    "Dalam proyek akhir ini, Google Gemini digunakan sebagai model generatif utama dalam "
    "arsitektur RAG. Gemini dipilih karena kemampuannya yang unggul dalam pemahaman dan "
    "generasi bahasa Indonesia, ketersediaan API yang mudah diakses melalui Google AI Studio, "
    "serta dukungan terhadap konteks percakapan multi-turn yang memungkinkan interaksi yang "
    "lebih natural. Model Gemini Pro yang digunakan dalam DamayAI memiliki kemampuan reasoning "
    "yang kuat dan mampu menghasilkan respons yang koheren berdasarkan konteks yang diberikan "
    "oleh hasil retrieval."
)

# 2.2.3
add_heading_left("2.2.3  Retrieval-Augmented Generation (RAG)")
add_body(
    "Retrieval-Augmented Generation (RAG) adalah pendekatan yang menggabungkan mekanisme "
    "retrieval (pencarian informasi) dengan kemampuan generatif LLM untuk menghasilkan "
    "jawaban yang lebih akurat dan terverifikasi. Konsep RAG pertama kali diperkenalkan oleh "
    "Lewis et al. (2020) dan sejak itu telah menjadi salah satu paradigma paling berpengaruh "
    "dalam pengembangan sistem tanya-jawab berbasis AI."
)
add_body(
    "Dalam arsitektur RAG, ketika pengguna mengajukan pertanyaan, sistem terlebih dahulu "
    "melakukan pencarian semantik (semantic search) terhadap basis pengetahuan yang telah "
    "diindeks untuk menemukan dokumen atau potongan teks yang paling relevan. Hasil "
    "retrieval ini kemudian digabungkan dengan pertanyaan asli pengguna sebagai konteks "
    "tambahan (augmented context) yang diberikan kepada LLM. Dengan demikian, LLM tidak "
    "hanya mengandalkan pengetahuan parametriknya, tetapi juga mendapatkan akses terhadap "
    "informasi aktual dari sumber eksternal yang telah diverifikasi."
)
add_body(
    "Proses RAG secara umum terdiri dari tiga tahap utama, yaitu: (1) Indexing—pembuatan "
    "indeks vektor dari dokumen sumber melalui proses chunking, embedding, dan penyimpanan "
    "di vector database; (2) Retrieval—pencarian dokumen yang paling relevan berdasarkan "
    "kemiripan semantik antara query pengguna dan dokumen yang diindeks; dan (3) Generation—"
    "proses pembuatan jawaban oleh LLM berdasarkan konteks yang diperoleh dari hasil "
    "retrieval."
)

# 2.2.3.1
add_heading_left("2.2.3.1  Keunggulan RAG dibanding Fine-tuning")
add_body(
    "Dalam konteks penyesuaian LLM untuk domain spesifik, terdapat dua pendekatan utama: "
    "fine-tuning dan RAG. Fine-tuning melibatkan pelatihan ulang model pada dataset "
    "domain-specific untuk mengubah pengetahuan parametrik model, sedangkan RAG "
    "mempertahankan model dasar dan menyediakan pengetahuan domain melalui mekanisme "
    "retrieval."
)
add_body("RAG memiliki beberapa keunggulan signifikan dibanding fine-tuning dalam konteks pengembangan chatbot informasi sekolah, antara lain sebagai berikut.")
add_bullet("RAG memungkinkan pembaruan pengetahuan secara real-time tanpa perlu melatih ulang model—cukup dengan menambahkan atau memperbarui dokumen di knowledge base.")
add_bullet("RAG menyediakan transparansi yang lebih baik karena jawaban dapat ditelusuri kembali ke sumber dokumen spesifik.")
add_bullet("RAG lebih efisien secara komputasional karena tidak memerlukan proses pelatihan yang mahal.")
add_bullet("RAG mengurangi risiko halusinasi karena LLM diarahkan untuk menghasilkan jawaban berdasarkan konteks yang disediakan, bukan mengandalkan pengetahuan parametrik yang mungkin sudah usang atau tidak akurat.")

# 2.2.4
add_heading_left("2.2.4  Metode Pengumpulan Data untuk Sistem RAG")
add_body(
    "Dalam pengembangan sistem RAG, kualitas dan kelengkapan data yang digunakan sebagai "
    "knowledge base sangat menentukan kualitas jawaban yang dihasilkan. DamayAI Assistant "
    "menggunakan tiga metode pengumpulan data utama yang masing-masing memiliki karakteristik "
    "dan tingkat prioritas yang berbeda dalam mekanisme Tiered RAG."
)

# 2.2.4.1
add_heading_left("2.2.4.1  Web Crawling dan Web Scraping")
add_body(
    "Web crawling adalah proses otomatis untuk menavigasi dan mengunduh halaman-halaman web "
    "secara sistematis, dimulai dari seed URL tertentu. Web scraping merupakan teknik "
    "lanjutan dari web crawling yang tidak hanya mengunduh halaman web, tetapi juga "
    "mengekstraksi informasi spesifik dari konten HTML. Dalam konteks DamayAI, web scraping "
    "digunakan untuk mengekstraksi informasi dari website resmi SMKN 2 Indramayu "
    "(smkn2indramayu.sch.id), termasuk data profil sekolah, program keahlian, data staf dan "
    "guru, serta informasi kegiatan sekolah."
)
add_body(
    "Proses web scraping pada DamayAI dilakukan menggunakan library BeautifulSoup dan "
    "Selenium untuk menangani halaman web yang bersifat statis maupun dinamis. Data yang "
    "diperoleh melalui scraping kemudian dibersihkan (cleaning), dinormalisasi, dan disimpan "
    "ke MongoDB sebelum akhirnya dipecah menjadi chunk dan di-embed ke FAISS. Data scraping "
    "memiliki prioritas terendah dalam mekanisme Tiered RAG karena kemungkinan adanya "
    "informasi yang belum terverifikasi atau tidak terbaru."
)

# 2.2.4.2
add_heading_left("2.2.4.2  Pengunggahan Dokumen Manual")
add_body(
    "Pengunggahan dokumen manual (manual upload) merupakan metode pengumpulan data di mana "
    "administrator mengunggah dokumen resmi sekolah dalam format PDF, DOCX, atau PPTX ke "
    "dalam sistem. Metode ini menghasilkan data dengan tingkat kepercayaan yang lebih tinggi "
    "dibanding scraping karena dokumen yang diunggah umumnya merupakan dokumen resmi yang "
    "telah diverifikasi oleh pihak sekolah, seperti profil lengkap staf, pedoman sekolah, "
    "dan dokumen akreditasi."
)
add_body(
    "Proses pengunggahan dokumen manual pada DamayAI melibatkan ekstraksi teks dari berbagai "
    "format file menggunakan library seperti PyPDF2 untuk PDF, python-docx untuk DOCX, dan "
    "python-pptx untuk PPTX. Teks yang diekstraksi kemudian melalui proses chunking dan "
    "embedding yang sama dengan data dari sumber lain. Data manual memiliki prioritas "
    "menengah dalam mekanisme Tiered RAG—lebih tinggi dari scraping namun lebih rendah dari "
    "Memory Bank."
)

# 2.2.4.3
add_heading_left("2.2.4.3  Memory Bank (Pembelajaran Berbasis Percakapan)")
add_body(
    "Memory Bank adalah fitur yang memungkinkan administrator untuk memasukkan pasangan "
    "pertanyaan-jawaban (Q&A pairs) secara langsung ke dalam sistem. Berbeda dengan "
    "scraping dan upload manual yang mengandalkan proses retrieval semantik, Memory Bank "
    "memungkinkan kurasi jawaban yang paling tepat dan terverifikasi untuk "
    "pertanyaan-pertanyaan spesifik yang sering diajukan (frequently asked questions)."
)
add_body(
    "Memory Bank memiliki prioritas tertinggi dalam mekanisme Tiered RAG karena setiap "
    "pasangan Q&A telah dikurasi secara manual oleh administrator yang memahami konteks dan "
    "kebutuhan informasi sekolah. Ketika pengguna mengajukan pertanyaan yang memiliki "
    "kemiripan semantik tinggi dengan entri di Memory Bank, sistem akan mengutamakan "
    "jawaban dari sumber ini terlebih dahulu sebelum mencari di lapisan data lainnya. "
    "Pendekatan ini memastikan bahwa jawaban yang paling tepercaya dan relevan selalu "
    "diprioritaskan."
)

# 2.2.5
add_heading_left("2.2.5  Text Embedding dan Vector Database")
add_heading_left("2.2.5.1  Text Embedding")
add_body(
    "Text embedding adalah proses mengkonversi teks menjadi representasi vektor numerik "
    "berdimensi tinggi yang menangkap makna semantik dari teks tersebut. Model embedding "
    "memetakan teks ke dalam ruang vektor sedemikian rupa sehingga teks dengan makna yang "
    "serupa berada berdekatan dalam ruang vektor tersebut. Teknik ini menjadi fondasi dari "
    "pencarian semantik (semantic search) yang digunakan dalam sistem RAG."
)
add_body(
    "Dalam DamayAI, digunakan model embedding dari Google (models/embedding-001) yang "
    "menghasilkan vektor berdimensi 768 untuk setiap input teks. Model ini dipilih karena "
    "kompatibilitasnya dengan ekosistem Google AI, kualitas embedding yang baik untuk bahasa "
    "Indonesia, dan efisiensi komputasional yang memadai untuk skala data sekolah. Proses "
    "embedding dilakukan pada setiap chunk teks yang dihasilkan dari proses chunking, dan "
    "vektor-vektor hasil embedding disimpan dalam FAISS untuk pencarian kemiripan yang cepat."
)

add_heading_left("2.2.5.2  FAISS (Facebook AI Similarity Search)")
add_body(
    "FAISS (Facebook AI Similarity Search) adalah library open-source yang dikembangkan oleh "
    "Facebook AI Research untuk pencarian kemiripan (similarity search) dan pengelompokan "
    "(clustering) vektor berdimensi tinggi secara efisien. FAISS mendukung berbagai metode "
    "indexing yang memungkinkan pencarian miliaran vektor dengan latensi yang sangat rendah."
)
add_body(
    "Dalam arsitektur DamayAI, FAISS berfungsi sebagai vector store yang menyimpan seluruh "
    "embedding dari dokumen yang telah diproses. Ketika pengguna mengajukan pertanyaan, "
    "query tersebut di-embed menggunakan model yang sama, dan FAISS melakukan pencarian "
    "nearest neighbor untuk menemukan chunk-chunk yang paling mirip secara semantik. DamayAI "
    "menggunakan IndexFlatL2 yang melakukan pencarian exhaustif berbasis jarak Euclidean "
    "(L2), yang cocok untuk skala data sekolah yang tidak terlalu besar namun memerlukan "
    "akurasi pencarian yang tinggi."
)

# 2.2.6
add_heading_left("2.2.6  Text Chunking (Pemecahan Teks)")
add_body(
    "Text chunking adalah proses memecah dokumen panjang menjadi potongan-potongan teks "
    "yang lebih kecil (chunk) sebelum dilakukan embedding. Proses ini diperlukan karena "
    "model embedding memiliki batasan panjang input, dan chunk yang lebih kecil menghasilkan "
    "pencarian semantik yang lebih akurat karena setiap chunk mewakili topik atau konsep "
    "yang lebih fokus."
)
add_body(
    "DamayAI menggunakan strategi fixed-size chunking dengan overlap, di mana teks dipecah "
    "menjadi chunk berukuran sekitar 500 karakter dengan overlap 100 karakter antar chunk. "
    "Overlap ini memastikan bahwa konteks yang melintasi batas chunk tidak hilang. Selain "
    "itu, dilakukan juga pembersihan teks (text cleaning) yang mencakup penghapusan karakter "
    "tidak perlu, normalisasi whitespace, dan penanganan karakter khusus sebelum proses "
    "chunking dilakukan."
)

# 2.2.7
add_heading_left("2.2.7  MongoDB sebagai Basis Data Dokumen")
add_body(
    "MongoDB adalah sistem manajemen basis data NoSQL berorientasi dokumen yang menyimpan "
    "data dalam format BSON (Binary JSON). MongoDB dipilih dalam pengembangan DamayAI karena "
    "kemampuannya dalam menyimpan dan mengelola data semi-terstruktur dan tidak terstruktur "
    "yang fleksibel, sesuai dengan karakteristik data sekolah yang beragam format dan "
    "skemanya."
)
add_body(
    "Dalam arsitektur DamayAI, MongoDB berfungsi sebagai basis data utama yang menyimpan "
    "metadata dokumen, log percakapan, data pengguna administrator, dan data Memory Bank. "
    "MongoDB juga menyimpan data mentah hasil scraping dan upload dokumen sebelum diproses "
    "lebih lanjut. Keunggulan MongoDB dalam hal fleksibilitas skema dan kemampuan horizontal "
    "scaling menjadikannya pilihan yang tepat untuk sistem yang memerlukan penyimpanan data "
    "dengan format yang bervariasi."
)

# 2.2.8
add_heading_left("2.2.8  Flask sebagai Framework Backend")
add_body(
    "Flask adalah micro-framework Python untuk pengembangan aplikasi web yang bersifat ringan "
    "dan modular. Flask dipilih sebagai framework backend DamayAI karena kesederhanaannya, "
    "fleksibilitas dalam pengembangan API, serta dukungan ekosistem Python yang luas untuk "
    "library-library AI dan NLP."
)
add_body(
    "Dalam arsitektur DamayAI, Flask menangani seluruh logika backend, termasuk routing API "
    "untuk chatbot, autentikasi administrator, pengelolaan knowledge base (scraping dan "
    "upload manual), pengelolaan Memory Bank, dan penyajian halaman web. Flask juga berperan "
    "sebagai penghubung antara komponen-komponen sistem seperti MongoDB, FAISS, dan Google "
    "Gemini API. Penggunaan Flask memungkinkan arsitektur yang bersih dan mudah di-maintain, "
    "dengan pemisahan yang jelas antara logika bisnis, data access, dan presentation layer."
)

# 2.2.9
add_heading_left("2.2.9  Aspek Keamanan dalam Sistem Berbasis LLM")
add_body(
    "Pengembangan sistem yang melibatkan LLM dan data institusional memerlukan perhatian "
    "khusus terhadap aspek keamanan. Dalam konteks DamayAI, terdapat beberapa aspek "
    "keamanan yang diperhatikan, yaitu sebagai berikut."
)
add_bullet("Keamanan akses—hanya administrator yang terautentikasi yang dapat mengakses panel pengelolaan knowledge base dan Memory Bank.")
add_bullet("Validasi input—mencegah injection attack melalui sanitasi input pengguna.")
add_bullet("Proteksi data—memastikan data internal sekolah tidak bocor melalui respons chatbot yang tidak semestinya.")
add_bullet("Rate limiting—mencegah penyalahgunaan sistem melalui pembatasan jumlah request.")
add_body(
    "Selain itu, aspek keamanan juga mencakup pencegahan prompt injection, di mana pengguna "
    "berupaya memanipulasi prompt sistem untuk mengakses informasi yang tidak diizinkan atau "
    "mengubah perilaku chatbot. DamayAI menerapkan sanitasi input dan system prompt yang "
    "ketat untuk meminimalkan risiko serangan ini. Kombinasi mekanisme keamanan ini "
    "memastikan bahwa sistem dapat dioperasikan secara aman dalam lingkungan institusi "
    "pendidikan."
)

# 2.3 Kerangka Berpikir
add_heading_left("2.3  Kerangka Berpikir")
add_body(
    "Kerangka berpikir dalam penelitian ini dibangun berdasarkan identifikasi permasalahan "
    "aksesibilitas informasi di SMKN 2 Indramayu yang kemudian diformulasikan menjadi solusi "
    "berbasis teknologi kecerdasan buatan. Alur berpikir dimulai dari identifikasi kondisi "
    "eksisting layanan informasi sekolah yang masih konvensional, kemudian dilakukan analisis "
    "kebutuhan pengguna (siswa, guru, staf TU), hingga perancangan solusi berupa chatbot "
    "asisten virtual yang mengimplementasikan pendekatan Tiered RAG. Secara sistematis, "
    "kerangka berpikir dapat diuraikan dalam tahapan-tahapan berikut."
)
add_numbered_item(1, "Identifikasi masalah: informasi tersebar, respons manual lambat, AI publik tidak memiliki data internal sekolah.")
add_numbered_item(2, "Analisis kebutuhan: sistem terpusat, real-time, berbasis data terverifikasi, dapat diakses kapan saja.")
add_numbered_item(3, "Perancangan solusi: chatbot berbasis web dengan arsitektur Tiered RAG menggunakan tiga lapisan knowledge base (Memory Bank, Data Manual, Data Scraping).")
add_numbered_item(4, "Implementasi: pengembangan sistem menggunakan Flask, MongoDB, FAISS, dan Google Gemini AI.")
add_numbered_item(5, "Pengujian: evaluasi fungsionalitas, akurasi jawaban, dan performa sistem melalui black box testing dan pengujian respons.")

doc.add_paragraph()
add_image_placeholder("Gambar 2.6  Kerangka Berpikir Penelitian")

# 2.4 Deskripsi Permasalahan
add_heading_left("2.4  Deskripsi Permasalahan")
add_body(
    "Berdasarkan identifikasi permasalahan yang telah diuraikan pada Bab I, deskripsi "
    "permasalahan dalam proyek akhir ini dapat dirinci sebagai berikut. SMKN 2 Indramayu "
    "sebagai institusi pendidikan yang besar dengan delapan program keahlian menghadapi "
    "tantangan signifikan dalam pengelolaan dan penyajian informasi kepada pemangku "
    "kepentingan, terutama siswa baru dan calon peserta didik."
)
add_body(
    "Permasalahan pertama berkaitan dengan fragmentasi sumber informasi. Saat ini, informasi "
    "sekolah tersebar di berbagai kanal yang tidak terintegrasi, mulai dari mading fisik yang "
    "hanya dapat diakses di lingkungan sekolah, grup WhatsApp yang tidak terorganisir dan "
    "sulit dicari, hingga website resmi sekolah yang tidak selalu diperbarui secara berkala. "
    "Kondisi ini menyebabkan siswa baru, yang merupakan pengguna informasi terbesar, "
    "kesulitan menemukan jawaban atas pertanyaan mereka secara mandiri dan efisien."
)
add_body(
    "Permasalahan kedua berkaitan dengan ketidaksesuaian solusi AI publik. Meskipun asisten "
    "AI seperti ChatGPT telah tersedia secara luas, sistem-sistem ini tidak memiliki akses "
    "terhadap data internal SMKN 2 Indramayu. Akibatnya, jawaban yang diberikan seringkali "
    "tidak relevan, tidak akurat, atau bahkan menyesatkan ketika digunakan untuk menanyakan "
    "informasi spesifik tentang sekolah, seperti nama guru, jadwal kegiatan, atau prosedur "
    "PPDB."
)
add_body(
    "Permasalahan ketiga berkaitan dengan beban kerja staf Tata Usaha. Setiap hari, staf TU "
    "harus menjawab puluhan pertanyaan serupa dari siswa baru dan orang tua, mulai dari "
    "pertanyaan tentang jam operasional sekolah, syarat PPDB, hingga lokasi ruang guru. "
    "Proses manual ini tidak hanya membebani staf TU tetapi juga menyebabkan keterlambatan "
    "respons dan ketidakkonsistenan informasi yang diberikan."
)

# Save checkpoint
doc.save(os.path.join(os.path.dirname(__file__), "DamayAI_PPA_Revised.docx"))
print("Checkpoint saved (Bab 2 complete).")
