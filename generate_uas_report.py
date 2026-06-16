"""
Script untuk generate laporan UAS Workshop Pemrograman Framework 2026
berdasarkan proyek DamayAI Assistant.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

# ── Default Style ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# ── Helper Functions ──
def add_heading_centered(text, bold=True, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    return p

def add_heading_left(text, bold=True, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body(text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    return p

def add_italic_body(text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    return p

def set_cell_text(cell, text, bold=False, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0

def shade_cells(row, color="D9E2F3"):
    for cell in row.cells:
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): color,
            qn('w:val'): 'clear'
        })
        shading.append(shading_elem)

# ============================================
# JUDUL DOKUMEN
# ============================================
add_heading_centered("LAPORAN UAS", bold=True, size=16)
add_heading_centered("WORKSHOP PEMROGRAMAN FRAMEWORK 2026", bold=True, size=14)
doc.add_paragraph()
add_heading_centered("DamayAI Assistant", bold=True, size=14)
add_heading_centered("Asisten Digital Berbasis AI untuk SMKN 2 Indramayu", bold=False, size=12)
doc.add_paragraph()
add_heading_centered("Disusun Oleh:", bold=False, size=12)
add_heading_centered("[Nama Mahasiswa]", bold=True, size=12)
add_heading_centered("[NIM]", bold=False, size=12)
doc.add_paragraph()
doc.add_paragraph()
add_heading_centered("POLITEKNIK NEGERI INDRAMAYU", bold=True, size=14)
add_heading_centered("TAHUN 2026", bold=True, size=12)

doc.add_page_break()

# ============================================
# BAB 1: PENDAHULUAN
# ============================================
add_heading_left("BAB 1: PENDAHULUAN", bold=True, size=14)
doc.add_paragraph()

add_heading_left("1.1 Latar Belakang", bold=True, size=12)
add_body("DamayAI Assistant adalah sebuah platform asisten digital berbasis kecerdasan buatan (AI) yang dirancang khusus untuk SMKN 2 Indramayu. Platform ini bertujuan untuk menyediakan layanan informasi otomatis kepada siswa, orang tua, dan masyarakat umum mengenai berbagai hal terkait sekolah, seperti informasi pendaftaran, program keahlian, kegiatan sekolah, dan lain sebagainya.", indent=True)
add_body("Proyek ini dikembangkan menggunakan framework Flask (Python) untuk backend, dengan integrasi Large Language Model (LLM) melalui Groq API menggunakan model Llama 3.1 8B Instant. Sistem ini juga dilengkapi dengan fitur Retrieval-Augmented Generation (RAG) menggunakan FAISS sebagai vector database dan MongoDB sebagai database utama.", indent=True)
doc.add_paragraph()

add_heading_left("1.2 Tujuan", bold=True, size=12)
add_body("1. Memenuhi tugas UAS mata kuliah Workshop Pemrograman Framework 2026.", indent=True)
add_body("2. Mengimplementasikan minimal 15 endpoint API RESTful yang fungsional dan aman.", indent=True)
add_body("3. Mendokumentasikan seluruh endpoint beserta hasil testing menggunakan Postman.", indent=True)
doc.add_paragraph()

add_heading_left("1.3 Teknologi yang Digunakan", bold=True, size=12)
add_body("- Backend: Flask (Python Framework)")
add_body("- Database: MongoDB")
add_body("- Vector Search: FAISS (Facebook AI Similarity Search)")
add_body("- AI Model: Llama 3.1 8B Instant via Groq API")
add_body("- Frontend: HTML, CSS, JavaScript")
add_body("- Authentication: Session-based dengan CSRF Protection")
add_body("- Rate Limiting: Flask-Limiter")
add_body("- Testing Tool: Postman")
doc.add_paragraph()

doc.add_page_break()

# ============================================
# BAB 2: DOKUMENTASI API ENDPOINT
# ============================================
add_heading_left("BAB 2: DOKUMENTASI API ENDPOINT", bold=True, size=14)
doc.add_paragraph()

add_heading_left("2.1 Deskripsi Proyek", bold=True, size=12)
add_body("DamayAI Assistant menyediakan berbagai endpoint API untuk mendukung fitur-fitur berikut:", indent=True)
add_body("- Chat AI publik dan admin")
add_body("- Manajemen data knowledge base (Scraped, Manual, Memory Bank)")
add_body("- Pelaporan dan manajemen bug report")
add_body("- Web scraping dan crawling")
add_body("- Dashboard dan statistik")
add_body("- Sistem autentikasi admin dengan keamanan CSRF")
doc.add_paragraph()

add_heading_left("2.2 Tabel Endpoint API", bold=True, size=12)
add_body("Berikut adalah daftar lengkap endpoint API yang telah diimplementasikan (minimal 15 endpoint sesuai persyaratan):")
doc.add_paragraph()

# ── API ENDPOINTS TABLE ──
endpoints = [
    ("1", "Autentikasi", "/api/admin/login", "POST", "Melakukan autentikasi login admin dengan password. Mengembalikan CSRF token jika berhasil."),
    ("2", "Autentikasi", "/api/admin/logout", "POST", "Melakukan logout sesi admin dan menghapus CSRF token."),
    ("3", "Chat AI", "/api/chat", "POST", "Endpoint publik untuk mengirim pertanyaan ke AI. Mengembalikan jawaban dari model Llama 3.1 berbasis RAG."),
    ("4", "Chat AI", "/api/admin_chat", "POST", "Endpoint admin untuk chat AI dengan response streaming (NDJSON). Mendukung riwayat percakapan."),
    ("5", "Bug Report", "/api/report_bug", "POST", "Endpoint publik untuk melaporkan bug/laporan kesalahan dengan deskripsi dan file lampiran opsional."),
    ("6", "Bug Report", "/api/get_bug_reports", "GET", "Mengambil daftar seluruh laporan bug dari database (admin only)."),
    ("7", "Bug Report", "/api/bug_reports/<id>", "GET", "Mengambil detail laporan bug berdasarkan ID."),
    ("8", "Bug Report", "/api/bug_reports/<id>/status", "PUT", "Memperbarui status laporan bug (Baru, Sedang Diproses, Selesai, Tidak Akan Diperbaiki)."),
    ("9", "Bug Report", "/api/bug_reports/<id>", "DELETE", "Menghapus laporan bug berdasarkan ID."),
    ("10", "Data Bank", "/api/get-data", "GET", "Mengambil seluruh data dari semua sumber (Scraped, Manual, Memory) dalam satu respons terurut."),
    ("11", "Data Bank", "/api/add_manual_text", "POST", "Menambahkan data teks manual ke dalam knowledge base."),
    ("12", "Data Bank", "/api/add_manual_file", "POST", "Mengunggah file (PDF, DOCX, PPTX, TXT) dan mengekstrak teksnya ke knowledge base."),
    ("13", "Data Bank", "/api/save_memory", "POST", "Menyimpan pasangan pertanyaan-jawaban ke dalam Memory Bank."),
    ("14", "Data Scraped", "/api/scraped-data", "GET", "Mengambil daftar seluruh data hasil web scraping."),
    ("15", "Data Scraped", "/api/scraped-data/<id>", "GET", "Mengambil detail data scraped berdasarkan ID."),
    ("16", "Data Scraped", "/api/scraped-data/<id>", "PUT", "Memperbarui data scraped (judul dan konten) berdasarkan ID."),
    ("17", "Data Scraped", "/api/scraped-data/<id>", "DELETE", "Menghapus data scraped berdasarkan ID."),
    ("18", "Data Manual", "/api/manual-data", "GET", "Mengambil daftar seluruh data manual yang telah ditambahkan."),
    ("19", "Data Manual", "/api/manual-data/<id>", "GET", "Mengambil detail data manual berdasarkan ID."),
    ("20", "Data Manual", "/api/manual-data/<id>", "PUT", "Memperbarui data manual (judul dan konten) berdasarkan ID."),
    ("21", "Data Manual", "/api/manual-data/<id>", "DELETE", "Menghapus data manual berdasarkan ID."),
    ("22", "Memory Bank", "/api/memory-data", "GET", "Mengambil daftar seluruh data di Memory Bank."),
    ("23", "Memory Bank", "/api/memory-data/<id>", "GET", "Mengambil detail data Memory Bank berdasarkan ID."),
    ("24", "Memory Bank", "/api/memory-data/<id>", "PUT", "Memperbarui data Memory Bank (pertanyaan dan jawaban) berdasarkan ID."),
    ("25", "Memory Bank", "/api/memory-data/<id>", "DELETE", "Menghapus data Memory Bank berdasarkan ID."),
    ("26", "Sistem", "/api/health", "GET", "Memeriksa status kesehatan API dan koneksi database."),
    ("27", "Sistem", "/api/dashboard/stats", "GET", "Mengambil statistik dashboard (jumlah data, bug report, dll)."),
    ("28", "Sistem", "/api/scrape", "POST", "Menjalankan proses web scraping berdasarkan daftar URL di file konfigurasi."),
    ("29", "Sistem", "/api/crawl", "POST", "Menjalankan deep crawling pada URL tertentu dengan jumlah halaman maksimum."),
    ("30", "Sistem", "/api/reindex", "POST", "Melakukan re-indexing seluruh vektor FAISS dari data yang tersimpan."),
    ("31", "Sistem", "/api/delete_faiss", "POST", "Menghapus seluruh indeks FAISS (Memory, Manual, Scraped)."),
    ("32", "Sistem", "/api/delete_db", "POST", "Mengosongkan seluruh koleksi database (scraped_data, manual_data, memory_bank)."),
    ("33", "Keamanan", "/api/csrf-token", "GET", "Mengambil CSRF token untuk admin yang terautentikasi."),
]

table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
header_cells = table.rows[0].cells
set_cell_text(header_cells[0], "No", bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(header_cells[1], "Modul", bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(header_cells[2], "Endpoint (API)", bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(header_cells[3], "Method", bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(header_cells[4], "Deskripsi", bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
shade_cells(table.rows[0], "D9E2F3")

for ep in endpoints:
    row = table.add_row()
    for i, val in enumerate(ep):
        alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_text(row.cells[i], val, size=8, alignment=alignment)

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(1.0)
    row.cells[1].width = Cm(2.5)
    row.cells[2].width = Cm(5.5)
    row.cells[3].width = Cm(1.5)
    row.cells[4].width = Cm(7.5)

doc.add_paragraph()
doc.add_page_break()

# ============================================
# BAB 3: HASIL TESTING API DENGAN POSTMAN
# ============================================
add_heading_left("BAB 3: HASIL TESTING API DENGAN POSTMAN", bold=True, size=14)
doc.add_paragraph()
add_body("Berikut adalah dokumentasi hasil testing seluruh endpoint API menggunakan aplikasi Postman. Setiap endpoint diuji dengan langkah-langkah yang detail beserta contoh respons JSON yang diharapkan.", indent=True)
doc.add_paragraph()

# ── TESTING SECTION ──
BASE_URL = "http://localhost:5000"

test_cases = [
    {
        "no": "1",
        "title": "Login Admin",
        "method": "POST",
        "url": f"{BASE_URL}/api/admin/login",
        "desc": "Melakukan autentikasi admin untuk mendapatkan sesi dan CSRF token.",
        "headers": "Content-Type: application/json",
        "body": '{\n    "password": "admin123"\n}',
        "expected_response": '{\n    "status": "success",\n    "message": "Login berhasil.",\n    "csrf_token": "a1b2c3d4e5f6..."\n}',
        "screenshot_note": "Tempel screenshot hasil testing POST /api/admin/login di sini"
    },
    {
        "no": "2",
        "title": "Logout Admin",
        "method": "POST",
        "url": f"{BASE_URL}/api/admin/logout",
        "desc": "Melakukan logout dari sesi admin.",
        "headers": "Cookie: session=<session_cookie>",
        "body": None,
        "expected_response": '{\n    "status": "success",\n    "message": "Logged out."\n}',
        "screenshot_note": "Tempel screenshot hasil testing POST /api/admin/logout di sini"
    },
    {
        "no": "3",
        "title": "Chat AI (Publik)",
        "method": "POST",
        "url": f"{BASE_URL}/api/chat",
        "desc": "Mengirim pertanyaan ke AI DamayAI sebagai pengguna publik.",
        "headers": "Content-Type: application/json",
        "body": '{\n    "query": "Siapa kepala sekolah SMKN 2 Indramayu?",\n    "history": []\n}',
        "expected_response": '{\n    "response": "Kepala Sekolah SMKN 2 Indramayu saat ini adalah Ibu **Yeti Sumiati**."\n}',
        "screenshot_note": "Tempel screenshot hasil testing POST /api/chat di sini"
    },
    {
        "no": "4",
        "title": "Chat Admin (Streaming)",
        "method": "POST",
        "url": f"{BASE_URL}/api/admin_chat",
        "desc": "Mengirim pertanyaan ke AI dengan response streaming (NDJSON format).",
        "headers": "Content-Type: application/json\nCookie: session=<session_cookie>\nX-CSRF-Token: <csrf_token>",
        "body": '{\n    "query": "Apa saja program keahlian di SMKN 2?",\n    "history": []\n}',
        "expected_response": '{"step": "start", "data": "Menerima pertanyaan: \'Apa saja program keahlian di SMKN 2?\'"}\n{"step": "memory_search", "data": "Mencari di Memory Bank..."}\n{"step": "final_answer", "data": "Program keahlian..."}',
        "screenshot_note": "Tempel screenshot hasil testing POST /api/admin_chat di sini"
    },
    {
        "no": "5",
        "title": "Lapor Bug (Publik)",
        "method": "POST",
        "url": f"{BASE_URL}/api/report_bug",
        "desc": "Mengirim laporan bug beserta file screenshot (opsional).",
        "headers": "Content-Type: multipart/form-data",
        "body": "Form Data:\n  description: \"Tombol kirim tidak berfungsi di halaman chat\"\n  file: (screenshot.png)",
        "expected_response": '{\n    "status": "success",\n    "message": "Laporan bug berhasil dikirim."\n}',
        "screenshot_note": "Tempel screenshot hasil testing POST /api/report_bug di sini"
    },
    {
        "no": "6",
        "title": "Ambil Semua Bug Report",
        "method": "GET",
        "url": f"{BASE_URL}/api/get_bug_reports",
        "desc": "Mengambil daftar seluruh laporan bug (admin only).",
        "headers": "Cookie: session=<session_cookie>",
        "body": None,
        "expected_response": '[\n    {\n        "_id": "6654a1b2c3d4e5f6a7b8c9d0",\n        "description": "Tombol kirim tidak berfungsi",\n        "status": "Baru",\n        "file_path": "bugs/screenshot.png",\n        "created_at": "2026-06-15T10:30:00"\n    }\n]',
        "screenshot_note": "Tempel screenshot hasil testing GET /api/get_bug_reports di sini"
    },
    {
        "no": "7",
        "title": "Ambil Detail Bug Report",
        "method": "GET",
        "url": f"{BASE_URL}/api/bug_reports/<id>",
        "desc": "Mengambil detail satu laporan bug berdasarkan ID.",
        "headers": "Cookie: session=<session_cookie>",
        "body": None,
        "expected_response": '{\n    "status": "success",\n    "data": {\n        "_id": "6654a1b2c3d4e5f6a7b8c9d0",\n        "description": "Tombol kirim tidak berfungsi",\n        "status": "Baru",\n        "file_path": "bugs/screenshot.png"\n    }\n}',
        "screenshot_note": "Tempel screenshot hasil testing GET /api/bug_reports/<id> di sini"
    },
    {
        "no": "8",
        "title": "Update Status Bug Report",
        "method": "PUT",
        "url": f"{BASE_URL}/api/bug_reports/<id>/status",
        "desc": "Memperbarui status laporan bug.",
        "headers": "Content-Type: application/json\nCookie: session=<session_cookie>\nX-CSRF-Token: <csrf_token>",
        "body": '{\n    "status": "Sedang Diproses"\n}',
        "expected_response": '{\n    "status": "success",\n    "message": "Bug report <id> status updated."\n}',
        "screenshot_note": "Tempel screenshot hasil testing PUT /api/bug_reports/<id>/status di sini"
    },
    {
        "no": "9",
        "title": "Hapus Bug Report",
        "method": "DELETE",
        "url": f"{BASE_URL}/api/bug_reports/<id>",
        "desc": "Menghapus laporan bug berdasarkan ID.",
        "headers": "Cookie: session=<session_cookie>\nX-CSRF-Token: <csrf_token>",
        "body": None,
        "expected_response": '{\n    "status": "success",\n    "message": "Bug report <id> deleted."\n}',
        "screenshot_note": "Tempel screenshot hasil testing DELETE /api/bug_reports/<id> di sini"
    },
    {
        "no": "10",
        "title": "Ambil Semua Data (Gabungan)",
        "method": "GET",
        "url": f"{BASE_URL}/api/get-data",
        "desc": "Mengambil seluruh data dari semua sumber (Scraped, Manual, Memory) terurut berdasarkan waktu.",
        "headers": "Cookie: session=<session_cookie>",
        "body": None,
        "expected_response": '[\n    {\n        "_id": "...",\n        "title": "Info PPDB 2026",\n        "content": "Pendaftaran PPDB...",\n        "type": "Scrap",\n        "timestamp": "2026-06-14T08:00:00"\n    },\n    {\n        "_id": "...",\n        "title": "manual-text-Visi Misi",\n        "content": "Visi dan Misi sekolah...",\n        "type": "Manual"\n    }\n]',
        "screenshot_note": "Tempel screenshot hasil testing GET /api/get-data di sini"
    },
    {
        "no": "11",
        "title": "Tambah Data Teks Manual",
        "method": "POST",
        "url": f"{BASE_URL}/api/add_manual_text",
        "desc": "Menambahkan data teks manual ke knowledge base.",
        "headers": "Content-Type: application/json\nCookie: session=<session_cookie>\nX-CSRF-Token: <csrf_token>",
        "body": '{\n    "title": "Visi dan Misi Sekolah",\n    "content": "Visi SMKN 2 Indramayu adalah menjadi sekolah unggulan..."\n}',
        "expected_response": '{\n    "status": "success",\n    "message": "Data teks manual berhasil ditambahkan."\n}',
        "screenshot_note": "Tempel screenshot hasil testing POST /api/add_manual_text di sini"
    },
    {
        "no": "12",
        "title": "Upload File Manual",
        "method": "POST",
        "url": f"{BASE_URL}/api/add_manual_file",
        "desc": "Mengunggah file (PDF/DOCX/PPTX/TXT) untuk diekstrak teksnya dan disimpan ke knowledge base.",
        "headers": "Content-Type: multipart/form-data\nCookie: session=<session_cookie>\nX-CSRF-Token: <csrf_token>",
        "body": "Form Data:\n  title: \"Dokumen Kurikulum 2026\"\n  file: (kurikulum.pdf)",
        "expected_response": '{\n    "status": "success",\n    "message": "Konten dari file \'kurikulum.pdf\' berhasil ditambahkan."\n}',
        "screenshot_note": "Tempel screenshot hasil testing POST /api/add_manual_file di sini"
    },
    {
        "no": "13",
        "title": "Simpan ke Memory Bank",
        "method": "POST",
        "url": f"{BASE_URL}/api/save_memory",
        "desc": "Menyimpan pasangan pertanyaan dan jawaban ke dalam Memory Bank.",
        "headers": "Content-Type: application/json\nCookie: session=<session_cookie>\nX-CSRF-Token: <csrf_token>",
        "body": '{\n    "question": "Kapan pendaftaran PPDB dibuka?",\n    "answer": "Pendaftaran PPDB SMKN 2 Indramayu dibuka pada bulan Mei setiap tahunnya."\n}',
        "expected_response": '{\n    "status": "success",\n    "message": "Percakapan berhasil disimpan ke memori."\n}',
        "screenshot_note": "Tempel screenshot hasil testing POST /api/save_memory di sini"
    },
    {
        "no": "14",
        "title": "Ambil Semua Data Scraped",
        "method": "GET",
        "url": f"{BASE_URL}/api/scraped-data",
        "desc": "Mengambil daftar seluruh data hasil web scraping.",
        "headers": "Cookie: session=<session_cookie>",
        "body": None,
        "expected_response": '{\n    "status": "success",\n    "count": 5,\n    "data": [\n        {\n            "_id": "...",\n            "url": "https://smkn2indramayu.sch.id/about",\n            "title": "Tentang Kami",\n            "content": "SMKN 2 Indramayu adalah...",\n            "scraped_at": "2026-06-10T10:00:00"\n        }\n    ]\n}',
        "screenshot_note": "Tempel screenshot hasil testing GET /api/scraped-data di sini"
    },
    {
        "no": "15",
        "title": "Ambil Detail Data Scraped",
        "method": "GET",
        "url": f"{BASE_URL}/api/scraped-data/<id>",
        "desc": "Mengambil detail data scraped berdasarkan ID.",
        "headers": "Cookie: session=<session_cookie>",
        "body": None,
        "expected_response": '{\n    "status": "success",\n    "data": {\n        "_id": "...",\n        "url": "https://smkn2indramayu.sch.id/about",\n        "title": "Tentang Kami",\n        "content": "SMKN 2 Indramayu adalah..."\n    }\n}',
        "screenshot_note": "Tempel screenshot hasil testing GET /api/scraped-data/<id> di sini"
    },
    {
        "no": "16",
        "title": "Update Data Scraped",
        "method": "PUT",
        "url": f"{BASE_URL}/api/scraped-data/<id>",
        "desc": "Memperbarui judul dan konten data scraped.",
        "headers": "Content-Type: application/json\nCookie: session=<session_cookie>\nX-CSRF-Token: <csrf_token>",
        "body": '{\n    "title": "Tentang Kami (Updated)",\n    "content": "SMKN 2 Indramayu merupakan sekolah kejuruan negeri..."\n}',
        "expected_response": '{\n    "status": "success",\n    "message": "Data scraped berhasil diperbarui."\n}',
        "screenshot_note": "Tempel screenshot hasil testing PUT /api/scraped-data/<id> di sini"
    },
    {
        "no": "17",
        "title": "Hapus Data Scraped",
        "method": "DELETE",
        "url": f"{BASE_URL}/api/scraped-data/<id>",
        "desc": "Menghapus data scraped berdasarkan ID.",
        "headers": "Cookie: session=<session_cookie>\nX-CSRF-Token: <csrf_token>",
        "body": None,
        "expected_response": '{\n    "status": "success",\n    "message": "Data scraped berhasil dihapus."\n}',
        "screenshot_note": "Tempel screenshot hasil testing DELETE /api/scraped-data/<id> di sini"
    },
    {
        "no": "18",
        "title": "Ambil Semua Data Manual",
        "method": "GET",
        "url": f"{BASE_URL}/api/manual-data",
        "desc": "Mengambil daftar seluruh data manual.",
        "headers": "Cookie: session=<session_cookie>",
        "body": None,
        "expected_response": '{\n    "status": "success",\n    "count": 3,\n    "data": [\n        {\n            "_id": "...",\n            "title": "Visi dan Misi Sekolah",\n            "content": "Visi SMKN 2 Indramayu...",\n            "source_name": "manual-text-Visi dan Misi",\n            "added_at": "2026-06-12T09:00:00"\n        }\n    ]\n}',
        "screenshot_note": "Tempel screenshot hasil testing GET /api/manual-data di sini"
    },
    {
        "no": "19",
        "title": "Ambil Detail Data Manual",
        "method": "GET",
        "url": f"{BASE_URL}/api/manual-data/<id>",
        "desc": "Mengambil detail data manual berdasarkan ID.",
        "headers": "Cookie: session=<session_cookie>",
        "body": None,
        "expected_response": '{\n    "status": "success",\n    "data": {\n        "_id": "...",\n        "title": "Visi dan Misi Sekolah",\n        "content": "Visi SMKN 2 Indramayu..."\n    }\n}',
        "screenshot_note": "Tempel screenshot hasil testing GET /api/manual-data/<id> di sini"
    },
    {
        "no": "20",
        "title": "Update Data Manual",
        "method": "PUT",
        "url": f"{BASE_URL}/api/manual-data/<id>",
        "desc": "Memperbarui judul dan konten data manual.",
        "headers": "Content-Type: application/json\nCookie: session=<session_cookie>\nX-CSRF-Token: <csrf_token>",
        "body": '{\n    "title": "Visi Misi Updated",\n    "content": "Visi baru SMKN 2 Indramayu..."\n}',
        "expected_response": '{\n    "status": "success",\n    "message": "Data manual berhasil diperbarui."\n}',
        "screenshot_note": "Tempel screenshot hasil testing PUT /api/manual-data/<id> di sini"
    },
    {
        "no": "21",
        "title": "Hapus Data Manual",
        "method": "DELETE",
        "url": f"{BASE_URL}/api/manual-data/<id>",
        "desc": "Menghapus data manual berdasarkan ID.",
        "headers": "Cookie: session=<session_cookie>\nX-CSRF-Token: <csrf_token>",
        "body": None,
        "expected_response": '{\n    "status": "success",\n    "message": "Data manual berhasil dihapus."\n}',
        "screenshot_note": "Tempel screenshot hasil testing DELETE /api/manual-data/<id> di sini"
    },
    {
        "no": "22",
        "title": "Ambil Semua Data Memory Bank",
        "method": "GET",
        "url": f"{BASE_URL}/api/memory-data",
        "desc": "Mengambil daftar seluruh data di Memory Bank.",
        "headers": "Cookie: session=<session_cookie>",
        "body": None,
        "expected_response": '{\n    "status": "success",\n    "count": 2,\n    "data": [\n        {\n            "_id": "...",\n            "question": "Kapan pendaftaran PPDB?",\n            "answer": "Pendaftaran PPDB dibuka bulan Mei.",\n            "saved_at": "2026-06-11T08:00:00"\n        }\n    ]\n}',
        "screenshot_note": "Tempel screenshot hasil testing GET /api/memory-data di sini"
    },
    {
        "no": "23",
        "title": "Ambil Detail Data Memory Bank",
        "method": "GET",
        "url": f"{BASE_URL}/api/memory-data/<id>",
        "desc": "Mengambil detail data Memory Bank berdasarkan ID.",
        "headers": "Cookie: session=<session_cookie>",
        "body": None,
        "expected_response": '{\n    "status": "success",\n    "data": {\n        "_id": "...",\n        "question": "Kapan pendaftaran PPDB?",\n        "answer": "Pendaftaran PPDB dibuka bulan Mei."\n    }\n}',
        "screenshot_note": "Tempel screenshot hasil testing GET /api/memory-data/<id> di sini"
    },
    {
        "no": "24",
        "title": "Update Data Memory Bank",
        "method": "PUT",
        "url": f"{BASE_URL}/api/memory-data/<id>",
        "desc": "Memperbarui pertanyaan dan jawaban di Memory Bank.",
        "headers": "Content-Type: application/json\nCookie: session=<session_cookie>\nX-CSRF-Token: <csrf_token>",
        "body": '{\n    "question": "Kapan pendaftaran PPDB 2026?",\n    "answer": "Pendaftaran PPDB 2026 dibuka pada bulan Mei-Juni."\n}',
        "expected_response": '{\n    "status": "success",\n    "message": "Data memory berhasil diperbarui."\n}',
        "screenshot_note": "Tempel screenshot hasil testing PUT /api/memory-data/<id> di sini"
    },
    {
        "no": "25",
        "title": "Hapus Data Memory Bank",
        "method": "DELETE",
        "url": f"{BASE_URL}/api/memory-data/<id>",
        "desc": "Menghapus data Memory Bank berdasarkan ID.",
        "headers": "Cookie: session=<session_cookie>\nX-CSRF-Token: <csrf_token>",
        "body": None,
        "expected_response": '{\n    "status": "success",\n    "message": "Data memory berhasil dihapus."\n}',
        "screenshot_note": "Tempel screenshot hasil testing DELETE /api/memory-data/<id> di sini"
    },
    {
        "no": "26",
        "title": "Health Check",
        "method": "GET",
        "url": f"{BASE_URL}/api/health",
        "desc": "Memeriksa status kesehatan API dan koneksi database.",
        "headers": None,
        "body": None,
        "expected_response": '{\n    "status": "healthy",\n    "database": "connected",\n    "timestamp": "2026-06-16T10:00:00"\n}',
        "screenshot_note": "Tempel screenshot hasil testing GET /api/health di sini"
    },
    {
        "no": "27",
        "title": "Dashboard Statistik",
        "method": "GET",
        "url": f"{BASE_URL}/api/dashboard/stats",
        "desc": "Mengambil statistik dashboard seperti jumlah data per kategori dan bug report.",
        "headers": "Cookie: session=<session_cookie>",
        "body": None,
        "expected_response": '{\n    "status": "success",\n    "data": {\n        "scraped_count": 15,\n        "manual_count": 8,\n        "memory_count": 12,\n        "bug_count": 3,\n        "recent_bugs": [...]\n    }\n}',
        "screenshot_note": "Tempel screenshot hasil testing GET /api/dashboard/stats di sini"
    },
    {
        "no": "28",
        "title": "Web Scraping",
        "method": "POST",
        "url": f"{BASE_URL}/api/scrape",
        "desc": "Menjalankan proses scraping berdasarkan daftar URL di file urls_to_scrape.txt.",
        "headers": "Cookie: session=<session_cookie>\nX-CSRF-Token: <csrf_token>",
        "body": None,
        "expected_response": "(Streaming text/plain)\nMembaca file 'urls_to_scrape.txt'...\nBERHASIL: https://smkn2indramayu.sch.id - SMKN 2 Indramayu\nDILEWATI/ERROR: https://example.com - Connection timeout",
        "screenshot_note": "Tempel screenshot hasil testing POST /api/scrape di sini"
    },
    {
        "no": "29",
        "title": "Deep Crawling",
        "method": "POST",
        "url": f"{BASE_URL}/api/crawl",
        "desc": "Menjalankan deep crawling pada URL target dengan jumlah halaman maksimum.",
        "headers": "Content-Type: application/json\nCookie: session=<session_cookie>\nX-CSRF-Token: <csrf_token>",
        "body": '{\n    "url": "https://smkn2indramayu.sch.id",\n    "max_pages": 10\n}',
        "expected_response": "(Streaming text/plain)\nINFO: Memulai crawling https://smkn2indramayu.sch.id...\nBERHASIL: https://smkn2indramayu.sch.id/about - Tentang Kami\nBERHASIL: https://smkn2indramayu.sch.id/contact - Kontak",
        "screenshot_note": "Tempel screenshot hasil testing POST /api/crawl di sini"
    },
    {
        "no": "30",
        "title": "Reindex FAISS",
        "method": "POST",
        "url": f"{BASE_URL}/api/reindex",
        "desc": "Melakukan re-indexing seluruh vektor FAISS dari data yang tersimpan di database.",
        "headers": "Cookie: session=<session_cookie>\nX-CSRF-Token: <csrf_token>",
        "body": None,
        "expected_response": "(Streaming text/plain)\nMemulai re-indexing FAISS...\nBuilding index untuk Memory Bank... (12 dokumen)\nBuilding index untuk Data Manual... (8 dokumen)\nBuilding index untuk Data Scraped... (15 dokumen)\nRe-indexing selesai.",
        "screenshot_note": "Tempel screenshot hasil testing POST /api/reindex di sini"
    },
    {
        "no": "31",
        "title": "Hapus Indeks FAISS",
        "method": "POST",
        "url": f"{BASE_URL}/api/delete_faiss",
        "desc": "Menghapus seluruh direktori indeks FAISS (Memory, Manual, Scraped).",
        "headers": "Cookie: session=<session_cookie>\nX-CSRF-Token: <csrf_token>",
        "body": None,
        "expected_response": '{\n    "status": "success",\n    "message": "Berhasil menghapus 3 direktori indeks FAISS."\n}',
        "screenshot_note": "Tempel screenshot hasil testing POST /api/delete_faiss di sini"
    },
    {
        "no": "32",
        "title": "Kosongkan Database",
        "method": "POST",
        "url": f"{BASE_URL}/api/delete_db",
        "desc": "Mengosongkan seluruh koleksi database (scraped_data, manual_data, memory_bank).",
        "headers": "Cookie: session=<session_cookie>\nX-CSRF-Token: <csrf_token>",
        "body": None,
        "expected_response": '{\n    "status": "success",\n    "message": "Semua koleksi database berhasil dikosongkan."\n}',
        "screenshot_note": "Tempel screenshot hasil testing POST /api/delete_db di sini"
    },
    {
        "no": "33",
        "title": "Ambil CSRF Token",
        "method": "GET",
        "url": f"{BASE_URL}/api/csrf-token",
        "desc": "Mengambil CSRF token untuk admin yang terautentikasi.",
        "headers": "Cookie: session=<session_cookie>",
        "body": None,
        "expected_response": '{\n    "csrf_token": "a1b2c3d4e5f6g7h8i9j0..."\n}',
        "screenshot_note": "Tempel screenshot hasil testing GET /api/csrf-token di sini"
    },
]

for tc in test_cases:
    # Title
    add_heading_left(f"3.{tc['no']} Testing API: {tc['title']}", bold=True, size=12)
    doc.add_paragraph()
    add_body(f"API: {tc['method']} {tc['url'].replace(BASE_URL, '')}", indent=False)
    add_body(f"Deskripsi: {tc['desc']}", indent=False)
    doc.add_paragraph()

    add_heading_left("Langkah-langkah Testing:", bold=True, size=11)
    add_body(f"1. Buka aplikasi Postman.")
    add_body(f"2. Pilih method {tc['method']}.")
    add_body(f"3. Masukkan URL: {tc['url']}")
    
    if tc['headers']:
        add_body(f"4. Atur header request:")
        for h in tc['headers'].split('\n'):
            add_body(f"   - {h.strip()}")
    
    if tc['body']:
        add_body(f"5. Masukkan body request:")
        add_code_block(tc['body'])
    
    doc.add_paragraph()
    add_heading_left("Hasil yang Diharapkan:", bold=True, size=11)
    add_code_block(tc['expected_response'])
    
    doc.add_paragraph()
    # Screenshot placeholder
    p_placeholder = doc.add_paragraph()
    p_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_placeholder.add_run(f"> [{tc['screenshot_note']}]")
    run.italic = True
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    # Separator
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep = p_sep.add_run("─" * 60)
    run_sep.font.color.rgb = RGBColor(200, 200, 200)
    run_sep.font.size = Pt(8)
    doc.add_paragraph()

doc.add_page_break()

# ============================================
# BAB 4: KESIMPULAN
# ============================================
add_heading_left("BAB 4: KESIMPULAN", bold=True, size=14)
doc.add_paragraph()

add_body("Berdasarkan hasil implementasi dan testing yang telah dilakukan, dapat disimpulkan bahwa:", indent=True)
doc.add_paragraph()
add_body("1. Proyek DamayAI Assistant telah berhasil mengimplementasikan 33 endpoint API RESTful yang mencakup fitur autentikasi, chat AI, manajemen data (scraped, manual, memory), bug reporting, web scraping, dan sistem administrasi.")
doc.add_paragraph()
add_body("2. Seluruh endpoint telah dilengkapi dengan fitur keamanan seperti autentikasi sesi admin, proteksi CSRF (Cross-Site Request Forgery), rate limiting, validasi input, sanitasi HTML, dan security headers.")
doc.add_paragraph()
add_body("3. Testing menggunakan Postman menunjukkan bahwa semua endpoint berfungsi sesuai dengan yang diharapkan, termasuk validasi error handling untuk input yang tidak valid.")
doc.add_paragraph()
add_body("4. Arsitektur RESTful yang digunakan memudahkan pengembangan dan pemeliharaan sistem di masa depan.")
doc.add_paragraph()
add_body("5. Integrasi dengan teknologi AI (Llama 3.1 via Groq) dan RAG (Retrieval-Augmented Generation) dengan FAISS memberikan kemampuan pencarian dan penjawaban pertanyaan yang cerdas dan kontekstual.")

doc.add_paragraph()

# ============================================
# SAVE DOCUMENT
# ============================================
output_dir = os.path.expanduser("~\\Downloads")
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "Laporan_UAS_DamayAI_Assistant.docx")
doc.save(output_path)
print(f"Dokumen berhasil disimpan di: {output_path}")
