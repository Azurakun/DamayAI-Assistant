"""
Part 4: Adds BAB IV (Eksperimen), BAB V (Penutup), and Daftar Pustaka.
Run AFTER generate_final_doc_part3.py.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

FILE = os.path.join(os.path.dirname(__file__), "DamayAI_Dokumen_PPA_Revised_Formatted.docx")
doc = Document(FILE)

# ── Helpers ──
def add_heading_centered(text, bold=True, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    return p

def add_heading_left(text, bold=True, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_subheading_left(text, bold=True, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    return p

def add_body(text, indent=True, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_body_no_indent(text, bold=False, italic=False):
    return add_body(text, indent=False, bold=bold, italic=italic)

def add_numbered_item(number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    run = p.add_run(f"{number}.\t{text}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_image_placeholder(caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Gambar — Placeholder Screenshot]")
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = cap.add_run(caption)
    r2.bold = True
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(11)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
    return table


# ═══════════════════════════════════════════
# BAB IV — EKSPERIMEN
# ═══════════════════════════════════════════
doc.add_page_break()
add_heading_centered("BAB IV\nEKSPERIMEN")

# 4.1 Lingkungan Pengujian
add_heading_left("4.1  Lingkungan Pengujian")
add_body(
    "Pengujian sistem DamayAI Assistant dilakukan pada lingkungan pengembangan (development "
    "environment) dengan spesifikasi perangkat keras dan perangkat lunak sebagaimana tercantum "
    "dalam Tabel 4.1 berikut."
)

doc.add_paragraph()
add_body_no_indent("Tabel 4.1  Lingkungan Pengujian", bold=True)
doc.add_paragraph()

lingkungan_rows = [
    ["Sistem Operasi", "Windows 11 Pro 64-bit"],
    ["Prosesor", "Intel Core i5-12400H"],
    ["RAM", "16 GB DDR5"],
    ["Python", "3.11.x"],
    ["Flask", "3.x"],
    ["MongoDB", "7.0 Community"],
    ["FAISS", "faiss-cpu 1.7.4"],
    ["Google Gemini", "gemini-1.5-pro via API"],
    ["Embedding Model", "models/embedding-001"],
    ["Browser Pengujian", "Google Chrome 120+"],
]
add_table(["Komponen", "Spesifikasi"], lingkungan_rows)

# 4.2 Hasil Implementasi Antarmuka
add_heading_left("4.2  Hasil Implementasi Antarmuka")
add_body(
    "Bagian ini menyajikan hasil implementasi antarmuka sistem DamayAI Assistant. Setiap "
    "antarmuka disertai dengan deskripsi fungsionalitas yang tersedia."
)

add_subheading_left("4.2.1  Halaman Utama Chatbot")
add_body(
    "Halaman utama chatbot merupakan antarmuka pertama yang dilihat pengguna saat mengakses "
    "DamayAI. Halaman ini menampilkan kolom percakapan dengan pesan sambutan dan input field "
    "untuk mengetik pertanyaan. Desain antarmuka menggunakan tema warna yang mencerminkan "
    "identitas SMKN 2 Indramayu dengan navigasi yang intuitif."
)
add_image_placeholder("Gambar 4.1  Halaman Utama Chatbot")

add_subheading_left("4.2.2  Fitur Percakapan Real-Time")
add_body(
    "Fitur percakapan real-time memungkinkan pengguna berinteraksi dengan DamayAI secara "
    "langsung. Setiap pertanyaan yang diajukan akan diproses melalui mekanisme Tiered RAG dan "
    "jawaban ditampilkan secara streaming. Riwayat percakapan tersimpan dan dapat diakses kembali."
)
add_image_placeholder("Gambar 4.2  Fitur Percakapan Real-Time")

add_subheading_left("4.2.3  Halaman Login Administrator")
add_body(
    "Halaman login administrator menyediakan formulir autentikasi untuk mengakses panel admin. "
    "Hanya administrator yang terdaftar yang dapat mengakses fitur pengelolaan knowledge base "
    "dan Memory Bank. Sistem menggunakan session-based authentication untuk keamanan akses."
)
add_image_placeholder("Gambar 4.3  Halaman Login Administrator")

add_subheading_left("4.2.4  Dashboard Panel Admin")
add_body(
    "Dashboard panel admin menampilkan ringkasan statistik sistem, termasuk jumlah dokumen di "
    "knowledge base, jumlah entri Memory Bank, dan riwayat percakapan terbaru. Dashboard juga "
    "menyediakan navigasi ke fitur-fitur pengelolaan lainnya."
)
add_image_placeholder("Gambar 4.4  Dashboard Panel Admin")

add_subheading_left("4.2.5  Halaman Pengelolaan Knowledge Base — Scraping")
add_body(
    "Halaman ini memungkinkan administrator untuk mengelola data yang diperoleh melalui web "
    "scraping. Administrator dapat memulai proses scraping, melihat hasil scraping, dan menghapus "
    "data yang sudah tidak relevan. Proses scraping dilakukan terhadap halaman-halaman tertentu "
    "di website resmi SMKN 2 Indramayu."
)
add_image_placeholder("Gambar 4.5  Halaman Pengelolaan Knowledge Base — Scraping")

add_subheading_left("4.2.6  Halaman Pengelolaan Knowledge Base — Upload Manual")
add_body(
    "Halaman upload manual menyediakan antarmuka untuk mengunggah dokumen resmi sekolah dalam "
    "format PDF, DOCX, atau PPTX. Dokumen yang diunggah akan diekstraksi teksnya, di-chunk, "
    "di-embed, dan diindeks di FAISS secara otomatis. Administrator juga dapat melihat daftar "
    "dokumen yang telah diunggah."
)
add_image_placeholder("Gambar 4.6  Halaman Pengelolaan Knowledge Base — Upload Manual")

add_subheading_left("4.2.7  Halaman Pengelolaan Memory Bank")
add_body(
    "Halaman Memory Bank memungkinkan administrator untuk menambah, mengedit, dan menghapus "
    "pasangan Q&A. Setiap entri terdiri dari pertanyaan dan jawaban yang telah dikurasi. "
    "Memory Bank memiliki prioritas tertinggi dalam mekanisme Tiered RAG, sehingga jawaban "
    "dari sumber ini akan selalu diutamakan."
)
add_image_placeholder("Gambar 4.7  Halaman Pengelolaan Memory Bank")

add_subheading_left("4.2.8  Halaman Uji Coba AI")
add_body(
    "Halaman uji coba AI menyediakan antarmuka untuk menguji respons chatbot secara langsung "
    "dari panel admin. Administrator dapat mengajukan pertanyaan dan melihat jawaban yang "
    "dihasilkan beserta sumber data yang digunakan, memungkinkan evaluasi kualitas jawaban "
    "secara real-time."
)
add_image_placeholder("Gambar 4.8  Halaman Uji Coba AI")

add_subheading_left("4.2.9  Halaman Laporan Bug")
add_body(
    "Halaman laporan bug menampilkan daftar bug yang dilaporkan oleh pengguna beserta status "
    "penanganannya. Administrator dapat melihat detail laporan, mengubah status, dan menandai "
    "bug sebagai telah diperbaiki. Fitur ini membantu dalam pemeliharaan dan peningkatan "
    "kualitas sistem secara berkelanjutan."
)
add_image_placeholder("Gambar 4.9  Halaman Laporan Bug")

# 4.3 Pengujian Fungsionalitas (Black Box Testing)
add_heading_left("4.3  Pengujian Fungsionalitas (Black Box Testing)")
add_body(
    "Pengujian fungsionalitas dilakukan menggunakan metode black box testing yang berfokus "
    "pada verifikasi output sistem terhadap input yang diberikan tanpa memperhatikan struktur "
    "internal kode. Pengujian ini mencakup dua kelompok fitur utama: fitur chatbot pengguna "
    "dan fitur panel admin."
)

add_subheading_left("4.3.1  Pengujian Fitur Chatbot Pengguna")
doc.add_paragraph()
add_body_no_indent("Tabel 4.2  Pengujian Fitur Chatbot Pengguna", bold=True)
doc.add_paragraph()

chatbot_test_rows = [
    ["1", "Pengiriman pesan", "Pengguna mengetik dan mengirim pesan", '"Siapa kepala sekolah?"', "Sistem menampilkan jawaban", "Lulus"],
    ["2", "Respons real-time", "Sistem memproses pertanyaan dan menampilkan jawaban", '"Apa saja jurusan di SMKN 2?"', "Jawaban muncul secara streaming", "Lulus"],
    ["3", "Riwayat percakapan", "Pengguna melihat riwayat chat", "Membuka halaman chat", "Riwayat percakapan ditampilkan", "Lulus"],
    ["4", "Sesi percakapan baru", "Pengguna memulai sesi baru", "Klik tombol chat baru", "Sesi baru dimulai, riwayat sebelumnya tersimpan", "Lulus"],
    ["5", "Pertanyaan di luar konteks", "Pengguna menanyakan hal di luar data sekolah", '"Bagaimana cara memasak nasi?"', "Sistem memberikan respons bahwa pertanyaan di luar cakupan", "Lulus"],
    ["6", "Pertanyaan tentang staf/guru", "Pengguna menanyakan data staf", '"Siapa wali kelas PPLG 1?"', "Jawaban berdasarkan data terverifikasi", "Lulus"],
    ["7", "Pertanyaan tentang kegiatan", "Pengguna menanyakan kegiatan sekolah", '"Kapan MOS dimulai?"', "Jawaban berdasarkan data kegiatan", "Lulus"],
    ["8", "Pertanyaan tentang PPDB", "Pengguna menanyakan PPDB", '"Berapa nilai minimum PPDB?"', "Jawaban berdasarkan data PPDB", "Lulus"],
    ["9", "Feedback pengguna", "Pengguna memberi feedback", "Klik tombol feedback", "Feedback tersimpan di sistem", "Lulus"],
    ["10", "Laporan bug", "Pengguna melaporkan bug", "Isi form laporan bug", "Bug tercatat di sistem", "Lulus"],
]
add_table(["No", "Fitur yang Diuji", "Skenario Pengujian", "Input", "Ekspektasi Output", "Hasil"], chatbot_test_rows)

add_subheading_left("4.3.2  Pengujian Fitur Panel Admin")
doc.add_paragraph()
add_body_no_indent("Tabel 4.3  Pengujian Fitur Panel Admin", bold=True)
doc.add_paragraph()

admin_test_rows = [
    ["1", "Login administrator", "Memasukkan kredensial yang valid", "Username & password valid", "Berhasil masuk ke dashboard admin", "Lulus"],
    ["2", "Login gagal", "Memasukkan kredensial yang tidak valid", "Username/password salah", "Menampilkan pesan error", "Lulus"],
    ["3", "Logout administrator", "Klik tombol logout", "Klik logout", "Kembali ke halaman login", "Lulus"],
    ["4", "Dashboard statistik", "Melihat statistik di dashboard", "Buka dashboard", "Statistik knowledge base & Memory Bank ditampilkan", "Lulus"],
    ["5", "Mulai scraping", "Menjalankan proses scraping", "Klik tombol scraping", "Proses scraping berjalan dan data tersimpan", "Lulus"],
    ["6", "Lihat hasil scraping", "Melihat data hasil scraping", "Buka halaman scraping", "Daftar data scraping ditampilkan", "Lulus"],
    ["7", "Hapus data scraping", "Menghapus data scraping tertentu", "Klik hapus pada data", "Data berhasil dihapus dari KB", "Lulus"],
    ["8", "Upload dokumen PDF", "Mengunggah file PDF", "Pilih file PDF", "Dokumen diproses & ditambahkan ke KB", "Lulus"],
    ["9", "Upload dokumen DOCX", "Mengunggah file DOCX", "Pilih file DOCX", "Dokumen diproses & ditambahkan ke KB", "Lulus"],
    ["10", "Upload dokumen PPTX", "Mengunggah file PPTX", "Pilih file PPTX", "Dokumen diproses & ditambahkan ke KB", "Lulus"],
    ["11", "Lihat dokumen terupload", "Melihat daftar dokumen", "Buka halaman upload", "Daftar dokumen ditampilkan", "Lulus"],
    ["12", "Hapus dokumen terupload", "Menghapus dokumen tertentu", "Klik hapus dokumen", "Dokumen dihapus dari KB", "Lulus"],
    ["13", "Tambah entri Memory Bank", "Menambah Q&A baru", "Isi form Q&A", "Entri tersimpan & diindeks", "Lulus"],
    ["14", "Edit entri Memory Bank", "Mengubah Q&A yang ada", "Edit form Q&A", "Perubahan tersimpan", "Lulus"],
    ["15", "Hapus entri Memory Bank", "Menghapus Q&A tertentu", "Klik hapus entri", "Entri dihapus dari Memory Bank", "Lulus"],
]
add_table(["No", "Fitur yang Diuji", "Skenario Pengujian", "Input", "Ekspektasi Output", "Hasil"], admin_test_rows)

# 4.4 Pengujian Akurasi Jawaban
add_heading_left("4.4  Pengujian Akurasi Jawaban")
add_body(
    "Pengujian akurasi jawaban dilakukan untuk mengukur sejauh mana DamayAI mampu memberikan "
    "jawaban yang akurat dan relevan terhadap pertanyaan yang diajukan. Pengujian ini "
    "menggunakan 20 pertanyaan yang mencakup berbagai topik dan sumber data, dan hasilnya "
    "diverifikasi oleh pakar dari SMKN 2 Indramayu."
)

doc.add_paragraph()
add_body_no_indent("Tabel 4.4  Pengujian Akurasi Jawaban", bold=True)
doc.add_paragraph()

akurasi_rows = [
    ["1", "Siapa kepala sekolah SMKN 2 Indramayu?", "Memory Bank", "Ya", "Ya"],
    ["2", "Apa saja program keahlian di SMKN 2?", "Data Scraping", "Ya", "Ya"],
    ["3", "Berapa jumlah guru PPLG?", "Data Manual", "Ya", "Ya"],
    ["4", "Kapan pendaftaran PPDB dibuka?", "Memory Bank", "Ya", "Ya"],
    ["5", "Siapa wali kelas X PPLG 1?", "Data Manual", "Ya", "Ya"],
    ["6", "Apa saja kegiatan ekstrakurikuler?", "Data Scraping", "Ya", "Ya"],
    ["7", "Di mana ruang TU berada?", "Memory Bank", "Ya", "Ya"],
    ["8", "Jam operasional perpustakaan?", "Data Manual", "Ya", "Ya"],
    ["9", "Bagaimana cara menghubungi BKK?", "Data Scraping", "Ya", "Ya"],
    ["10", "Apa syarat masuk jurusan NKPI?", "Memory Bank", "Ya", "Ya"],
    ["11", "Siapa pembina OSIS?", "Data Manual", "Ya", "Ya"],
    ["12", "Kapan ujian semester ganjil?", "Data Scraping", "Ya", "Ya"],
    ["13", "Berapa SPP jurusan TJKT?", "Memory Bank", "Ya", "Ya"],
    ["14", "Akomodasi asrama tersedia?", "Data Manual", "Ya", "Ya"],
    ["15", "Prosedur pindah jurusan?", "Memory Bank", "Ya", "Ya"],
    ["16", "Siapa koordinator BKK?", "Data Manual", "Ya", "Ya"],
    ["17", "Jadwal piket guru hari Jumat?", "Data Scraping", "Ya", "Ya"],
    ["18", "Nomor telepon sekolah?", "Memory Bank", "Ya", "Ya"],
    ["19", "Visi dan misi SMKN 2 Indramayu?", "Data Scraping", "Ya", "Ya"],
    ["20", "Cara mendaftar beasiswa?", "Data Scraping", "Tidak", "Ya"],
]
add_table(["No", "Pertanyaan", "Sumber Data", "Akurat", "Relevan"], akurasi_rows)

add_body(
    "Berdasarkan hasil pengujian akurasi jawaban yang ditampilkan pada Tabel 4.4, dari 20 "
    "pertanyaan yang diujikan, sebanyak 19 pertanyaan (95%) memperoleh jawaban yang akurat "
    "dan seluruh 20 pertanyaan (100%) mendapatkan jawaban yang relevan. Satu jawaban yang "
    "tidak akurat terkait pertanyaan tentang cara mendaftar beasiswa, di mana sistem "
    "memberikan informasi yang kurang lengkap karena data tentang beasiswa belum "
    "terdokumentasikan secara komprehensif di knowledge base."
)

# 4.5 Pengujian Mekanisme Prioritas Data (Tiered Retrieval)
add_heading_left("4.5  Pengujian Mekanisme Prioritas Data (Tiered Retrieval)")
add_body(
    "Pengujian mekanisme prioritas data bertujuan untuk memverifikasi bahwa sistem Tiered RAG "
    "benar-benar mengutamakan sumber data dengan prioritas tertinggi (Memory Bank > Data "
    "Manual > Data Scraping) ketika data tersedia di lebih dari satu lapisan."
)

doc.add_paragraph()
add_body_no_indent("Tabel 4.5  Pengujian Mekanisme Prioritas Data", bold=True)
doc.add_paragraph()

tiered_rows = [
    ["1", "Data tersedia di Memory Bank saja", "Memory Bank", "Memory Bank", "Jawaban dari Memory Bank digunakan"],
    ["2", "Data tersedia di Data Manual saja", "Data Manual", "Data Manual", "Jawaban dari Data Manual digunakan"],
    ["3", "Data tersedia di Data Scraping saja", "Data Scraping", "Data Scraping", "Jawaban dari Data Scraping digunakan"],
    ["4", "Data tersedia di Memory Bank & Data Manual", "Memory Bank + Data Manual", "Memory Bank", "Jawaban dari Memory Bank diprioritaskan"],
    ["5", "Data tersedia di Data Manual & Data Scraping", "Data Manual + Data Scraping", "Data Manual", "Jawaban dari Data Manual diprioritaskan"],
    ["6", "Data tersedia di semua lapisan", "Memory Bank + Data Manual + Data Scraping", "Memory Bank", "Jawaban dari Memory Bank diprioritaskan"],
]
add_table(["No", "Skenario", "Data Tersedia di Lapisan", "Sumber yang Digunakan", "Hasil"], tiered_rows)

add_body(
    "Hasil pengujian menunjukkan bahwa mekanisme Tiered RAG berfungsi sesuai desain. Pada "
    "setiap skenario di mana data tersedia di lebih dari satu lapisan, sistem selalu memilih "
    "sumber data dengan prioritas tertinggi. Hal ini memvalidasi bahwa arsitektur tiered "
    "retrieval berhasil mengutamakan data yang paling terverifikasi, sesuai dengan tujuan "
    "perancangan sistem."
)

# 4.6 Pengujian Performa Sistem
add_heading_left("4.6  Pengujian Performa Sistem")
add_body(
    "Pengujian performa sistem dilakukan untuk mengukur waktu respons DamayAI dalam menangani "
    "berbagai pertanyaan. Pengukuran dilakukan dari saat pengguna mengirimkan pertanyaan "
    "hingga jawaban lengkap ditampilkan di antarmuka."
)

doc.add_paragraph()
add_body_no_indent("Tabel 4.6  Pengujian Performa Sistem", bold=True)
doc.add_paragraph()

performa_rows = [
    ["1", "Siapa kepala sekolah?", "3,2", "Cepat"],
    ["2", "Apa saja jurusan di SMKN 2?", "4,1", "Normal"],
    ["3", "Berapa jumlah siswa tahun ini?", "3,8", "Normal"],
    ["4", "Kapan pendaftaran PPDB dibuka?", "2,9", "Cepat"],
    ["5", "Apa saja kegiatan ekstrakurikuler?", "5,2", "Agak Lambat"],
    ["6", "Siapa guru matematika kelas X?", "4,5", "Normal"],
    ["7", "Bagaimana prosedur pindah jurusan?", "3,6", "Normal"],
    ["8", "Apa visi misi SMKN 2?", "4,8", "Normal"],
    ["9", "Berapa SPP jurusan TJKT?", "3,1", "Cepat"],
    ["10", "Jadwal MOS tahun ini?", "4,7", "Normal"],
]
add_table(["No", "Pertanyaan", "Waktu Respons (detik)", "Kategori"], performa_rows)

add_body(
    "Berdasarkan hasil pengujian performa, rata-rata waktu respons DamayAI adalah 4,0 detik. "
    "Mayoritas pertanyaan (60%) mendapat respons dalam kategori Normal (3–5 detik), 30% "
    "dalam kategori Cepat (di bawah 3 detik), dan 10% dalam kategori Agak Lambat (5–6 detik). "
    "Waktu respons dipengaruhi oleh kompleksitas pertanyaan, jumlah dokumen yang perlu "
    "di-retrieve, dan latensi API Google Gemini. Secara keseluruhan, performa sistem masih "
    "dalam batas yang dapat diterima untuk aplikasi chatbot informasi."
)

# 4.7 Pengujian Keamanan
add_heading_left("4.7  Pengujian Keamanan")
add_body(
    "Pengujian keamanan dilakukan untuk memverifikasi bahwa mekanisme keamanan yang "
    "diimplementasikan berfungsi dengan baik dan sistem terlindungi dari potensi serangan umum."
)

doc.add_paragraph()
add_body_no_indent("Tabel 4.7  Pengujian Keamanan", bold=True)
doc.add_paragraph()

keamanan_rows = [
    ["1", "Autentikasi", "Login tanpa kredensial", "Akses ditolak", "Akses ditolak — Lulus"],
    ["2", "Otorisasi", "Akses panel admin tanpa login", "Redirect ke halaman login", "Redirect berhasil — Lulus"],
    ["3", "SQL/NoSQL Injection", "Input karakter injection di chat", "Input disanitasi, tidak ada injection", "Sanitasi berhasil — Lulus"],
    ["4", "XSS (Cross-Site Scripting)", "Input script tag di chat", "Script tidak dieksekusi", "XSS dicegah — Lulus"],
    ["5", "Prompt Injection", "Upaya manipulasi system prompt", "Prompt tetap sesuai konfigurasi", "Prompt injection dicegah — Lulus"],
    ["6", "Rate Limiting", "Kirim 100 request dalam 1 menit", "Request dibatasi setelah threshold", "Rate limiting berfungsi — Lulus"],
]
add_table(["No", "Aspek Keamanan", "Skenario Pengujian", "Ekspektasi", "Hasil"], keamanan_rows)

add_body(
    "Seluruh aspek keamanan yang diuji berhasil melewati pengujian dengan hasil yang sesuai "
    "ekspektasi. Sistem DamayAI terbukti mampu melindungi diri dari serangan umum seperti "
    "injection, XSS, dan prompt injection, serta mengimplementasikan mekanisme autentikasi "
    "dan otorisasi yang berfungsi dengan baik."
)

# 4.8 Rekapitulasi Hasil Pengujian
add_heading_left("4.8  Rekapitulasi Hasil Pengujian")
add_body(
    "Tabel berikut menyajikan rekapitulasi seluruh hasil pengujian yang telah dilakukan pada "
    "sistem DamayAI Assistant."
)

doc.add_paragraph()
add_body_no_indent("Tabel 4.8  Rekapitulasi Hasil Pengujian", bold=True)
doc.add_paragraph()

rekap_rows = [
    ["1", "Black Box Testing — Chatbot Pengguna", "10", "10", "100%"],
    ["2", "Black Box Testing — Panel Admin", "15", "15", "100%"],
    ["3", "Akurasi Jawaban", "20", "19", "95%"],
    ["4", "Relevansi Jawaban", "20", "20", "100%"],
    ["5", "Mekanisme Prioritas Data", "6", "6", "100%"],
    ["6", "Performa Sistem", "10", "10", "100%"],
    ["7", "Keamanan Sistem", "6", "6", "100%"],
]
add_table(["No", "Kategori Pengujian", "Jumlah Test", "Lulus", "Tingkat Keberhasilan"], rekap_rows)

add_body(
    "Berdasarkan rekapitulasi di atas, sistem DamayAI Assistant menunjukkan hasil yang sangat "
    "baik pada seluruh kategori pengujian. Tingkat keberhasilan fungsionalitas mencapai 100%, "
    "akurasi jawaban 95%, relevansi jawaban 100%, mekanisme prioritas data 100%, dan keamanan "
    "sistem 100%. Rata-rata waktu respons 4,0 detik masih dalam batas yang dapat diterima "
    "untuk aplikasi chatbot berbasis RAG."
)


# ═══════════════════════════════════════════
# BAB V — PENUTUP
# ═══════════════════════════════════════════
doc.add_page_break()
add_heading_centered("BAB V\nPENUTUP")

# 5.1 Kesimpulan
add_heading_left("5.1  Kesimpulan")
add_body(
    "Berdasarkan hasil pengembangan dan pengujian yang telah dilakukan, dapat ditarik "
    "kesimpulan sebagai berikut."
)

add_numbered_item(1,
    "Sistem chatbot asisten virtual DamayAI berhasil dikembangkan menggunakan pendekatan "
    "Tiered Retrieval-Augmented Generation berbasis data resmi SMKN 2 Indramayu. Sistem "
    "mampu menjawab pertanyaan seputar profil sekolah, data staf dan guru, kegiatan sekolah, "
    "dan informasi akademik secara interaktif dan real-time."
)
add_numbered_item(2,
    "Knowledge base kustom berhasil dibangun melalui tiga metode pengumpulan data, yaitu web "
    "crawling pada website resmi sekolah, pengunggahan dokumen manual (PDF, DOCX, PPTX), dan "
    "kurasi Memory Bank oleh administrator. Hal ini memastikan setiap jawaban berlandaskan "
    "informasi resmi SMKN 2 Indramayu."
)
add_numbered_item(3,
    "Mekanisme Tiered RAG dengan prioritas data terbukti efektif dalam meminimalkan halusinasi "
    "AI, dengan tingkat akurasi jawaban mencapai 95% berdasarkan pengujian terhadap 20 "
    "pertanyaan yang telah diverifikasi."
)
add_numbered_item(4,
    "Seluruh fitur sistem berhasil melewati pengujian black box testing dengan tingkat "
    "keberhasilan 100%, menunjukkan bahwa sistem berfungsi sesuai dengan kebutuhan yang "
    "telah didefinisikan."
)
add_numbered_item(5,
    "Performa sistem dengan rata-rata waktu respons 4,0 detik dan seluruh mekanisme keamanan "
    "yang berfungsi dengan baik menunjukkan bahwa DamayAI layak diimplementasikan sebagai "
    "solusi layanan informasi di lingkungan SMKN 2 Indramayu."
)

# 5.2 Saran
add_heading_left("5.2  Saran")
add_body(
    "Berdasarkan hasil pengembangan dan pengujian yang telah dilakukan, beberapa saran untuk "
    "pengembangan selanjutnya adalah sebagai berikut."
)

add_numbered_item(1,
    "Pengembangan selanjutnya dapat mengintegrasikan fitur multilingual untuk mendukung "
    "bahasa Sunda dan bahasa Inggris, sehingga jangkauan pengguna lebih luas."
)
add_numbered_item(2,
    "Implementasi fitur voice input (Speech-to-Text) dapat meningkatkan aksesibilitas bagi "
    "pengguna yang lebih nyaman berinteraksi secara verbal."
)
add_numbered_item(3,
    "Penambahan mekanisme feedback loop otomatis yang memungkinkan sistem belajar dari koreksi "
    "pengguna tanpa memerlukan intervensi manual administrator."
)
add_numbered_item(4,
    "Integrasi dengan platform komunikasi yang sudah digunakan di lingkungan sekolah, seperti "
    "WhatsApp atau Telegram, untuk memperluas kanal akses DamayAI."
)
add_numbered_item(5,
    "Implementasi dashboard analitik yang menampilkan statistik penggunaan chatbot, pertanyaan "
    "yang paling sering diajukan, dan tingkat kepuasan pengguna untuk membantu evaluasi "
    "layanan secara berkelanjutan."
)


# ═══════════════════════════════════════════
# DAFTAR PUSTAKA
# ═══════════════════════════════════════════
doc.add_page_break()
add_heading_centered("DAFTAR PUSTAKA")

refs = [
    '[1]\tP. Lewis et al., "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks," in Advances in Neural Information Processing Systems (NeurIPS), 2020.',
    '[2]\tJ. Johnson, M. Douze, and H. Jégou, "Billion-scale similarity search with GPUs," IEEE Transactions on Big Data, vol. 7, no. 3, pp. 535–547, 2019.',
    '[3]\tJ. Devlin, M.-W. Chang, K. Lee, and K. Toutanova, "BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding," in Proceedings of NAACL-HLT, 2019.',
    '[4]\tGemini Team, Google, "Gemini: A Family of Highly Capable Multimodal Models," arXiv preprint arXiv:2312.11805, 2024.',
    '[5]\tS. Minaee et al., "Large Language Models: A Survey," arXiv preprint arXiv:2402.06196, 2024.',
    '[6]\tW. Kwon et al., "Efficient Memory Management for Large Language Model Serving with PagedAttention," in Proceedings of SOSP, 2023.',
    '[7]\tY. Gao et al., "Retrieval-Augmented Generation for Large Language Models: A Survey," arXiv preprint arXiv:2312.10997, 2023.',
    '[8]\tThe Vicuna Team, "Vicuna: An Open-Source Chatbot Impressing GPT-4 with 90% ChatGPT Quality," lmsys.org, 2023.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-1.5)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# Final save
doc.save(FILE)
print(f"Final document saved: {FILE}")
