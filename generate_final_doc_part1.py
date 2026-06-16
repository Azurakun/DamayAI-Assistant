"""
Script to generate the revised DamayAI PPA document (Complete)
following proper academic report structure with Times New Roman formatting.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

# ── Default Style ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# ── Helper Functions ──
def add_heading_centered(text, level=1, bold=True, size=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    return p

def add_heading_left(text, level=2, bold=True, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_subheading_left(text, bold=True, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    return p

def add_body(text, indent=True, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_body_no_indent(text, bold=False, italic=False):
    return add_body(text, indent=False, bold=bold, italic=italic)

def add_bullet(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run(f"•  {text}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_numbered_item(number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    run = p.add_run(f"{number}.\t{text}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_blank_lines(n=1):
    for _ in range(n):
        doc.add_paragraph()

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
    return table

def add_image_placeholder(caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Gambar — Placeholder Screenshot]")
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = cap.add_run(caption)
    r2.bold = True
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(11)


# ═══════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════
add_blank_lines(2)
add_heading_centered("PROPOSAL", size=14)
add_blank_lines(1)
add_heading_centered("LAPORAN AKHIR", size=14)
add_blank_lines(1)
add_heading_centered(
    "IMPLEMENTASI CHATBOT ASISTEN VIRTUAL SMKN 2 INDRAMAYU\n"
    "MENGGUNAKAN METODE TIERED RETRIEVAL-AUGMENTED GENERATION\n"
    "DENGAN MEKANISME PRIORITAS DATA", size=14)
add_blank_lines(2)
add_heading_centered("Achmad Arditio Sumartono", size=12)
add_heading_centered("NRP. 3124510063", size=12)
add_blank_lines(2)
add_heading_centered("DOSEN PEMBIMBING", size=12)
add_blank_lines(1)
add_heading_centered("Prof. Dr. Arna Fariza, S.Kom., M.Kom.", size=12)
add_heading_centered("NIP. 197107081999032001", size=11)
add_blank_lines(1)
add_heading_centered("Fitrah Maharani Humaira, M.Kom.", size=12)
add_heading_centered("NIP.", size=11)
add_blank_lines(3)
add_heading_centered(
    "PROGRAM STUDI DIPLOMA TIGA\n"
    "TEKNIK INFORMATIKA\n"
    "DEPARTEMEN TEKNIK INFORMATIKA DAN KOMPUTER\n"
    "POLITEKNIK ELEKTRONIKA NEGERI SURABAYA\n"
    "2025", size=12)

doc.add_page_break()

# ═══════════════════════════════════════════
# DAFTAR ISI (placeholder)
# ═══════════════════════════════════════════
add_heading_centered("DAFTAR ISI", size=14)
add_body_no_indent("(Daftar Isi akan di-generate secara otomatis oleh Microsoft Word)")
doc.add_page_break()

add_heading_centered("DAFTAR GAMBAR", size=14)
add_body_no_indent("(Daftar Gambar akan di-generate secara otomatis oleh Microsoft Word)")
doc.add_page_break()

add_heading_centered("DAFTAR TABEL", size=14)
add_body_no_indent("(Daftar Tabel akan di-generate secara otomatis oleh Microsoft Word)")
doc.add_page_break()


# ═══════════════════════════════════════════
# BAB I — PENDAHULUAN
# ═══════════════════════════════════════════
add_heading_centered("BAB I\nPENDAHULUAN", size=14)

# 1.1 Latar Belakang
add_heading_left("1.1  Latar Belakang")

add_body(
    "SMK Negeri 2 Indramayu merupakan salah satu sekolah kejuruan negeri unggulan di "
    "Kabupaten Indramayu, Jawa Barat, yang menyelenggarakan berbagai program keahlian "
    "berorientasi industri, mulai dari NKPI, APHPI, TJKT, PPLG, Teknik Alat Berat, "
    "Kuliner, APAPL, hingga TPFL. Dengan jumlah peserta didik yang terus bertumbuh dari "
    "tahun ke tahun, sekolah ini memiliki ekosistem informasi yang sangat luas, mulai dari "
    "jadwal kegiatan, data staf dan guru, informasi PPDB, hingga rekrutmen kerja yang "
    "difasilitasi oleh Bursa Kerja Khusus (BKK). Tingginya volume informasi ini tentu "
    "menjadi tantangan tersendiri bagi pengelolaan layanan informasi sekolah yang selama "
    "ini masih dilakukan secara konvensional dan tidak terpusat."
)

add_body(
    "Perkembangan kecerdasan buatan, khususnya hadirnya Large Language Model (LLM) "
    "seperti Google Gemini, telah membuka peluang besar untuk mengotomasi layanan "
    "informasi semacam ini secara lebih efisien dan kontekstual. Meski begitu, pengelolaan "
    "informasi di SMKN 2 Indramayu saat ini masih tersebar di berbagai kanal yang tidak "
    "terpusat, seperti mading fisik, grup WhatsApp, hingga website yang tidak selalu "
    "diperbarui. Kondisi ini mempersulit siswa baru untuk menemukan informasi secara "
    "cepat dan mandiri, serta membebani staf Tata Usaha yang harus menjawab "
    "pertanyaan-pertanyaan serupa secara berulang setiap harinya."
)

add_body(
    "Penggunaan asisten AI publik seperti ChatGPT pun tidak dapat menjadi solusi yang "
    "tepat karena sistem tersebut tidak memiliki akses terhadap data internal sekolah, "
    "sehingga rentan memberikan informasi yang tidak akurat atau bahkan menyesatkan. "
    "Pendekatan Retrieval-Augmented Generation (RAG) telah terbukti mampu mengatasi "
    "permasalahan ini dengan menggabungkan kemampuan generatif LLM dan pencarian semantik "
    "berbasis vector store seperti FAISS, sehingga jawaban yang dihasilkan dapat didasarkan "
    "pada dokumen terverifikasi milik institusi."
)

add_body(
    "Berdasarkan kajian terhadap penelitian terdahulu, sistem chatbot berbasis RAG untuk "
    "institusi pendidikan telah banyak dikembangkan di lingkungan universitas luar negeri, "
    "namun belum ada yang secara spesifik menyasar konteks sekolah menengah kejuruan (SMK) "
    "di Indonesia, khususnya dengan pendekatan tiered retrieval tiga lapis yang membedakan "
    "antara Memory Bank, Data Manual, dan Data Scraping. Oleh karena itu, pada proyek "
    "akhir ini dikembangkanlah DamayAI Assistant, sebuah chatbot asisten virtual yang "
    "dibangun menggunakan pendekatan Tiered Retrieval-Augmented Generation berbasis data "
    "resmi SMKN 2 Indramayu, mencakup profil sekolah, data staf, kegiatan, dan informasi "
    "akademik, sehingga setiap jawaban dapat diprioritaskan berdasarkan sumber data yang "
    "paling tepercaya."
)

# 1.2 Identifikasi Permasalahan
add_heading_left("1.2  Identifikasi Permasalahan")

add_body(
    "Berdasarkan latar belakang yang telah diuraikan, terdapat tiga permasalahan utama "
    "yang diidentifikasi sebagai berikut."
)

add_numbered_item(1,
    "Informasi seputar SMKN 2 Indramayu, seperti data staf, jadwal kegiatan, profil "
    "sekolah, dan informasi PPDB, masih tersebar di berbagai kanal yang tidak terpusat, "
    "termasuk mading fisik, grup WhatsApp, dan website yang tidak selalu diperbarui, "
    "sehingga siswa baru kesulitan mengakses informasi secara cepat dan mandiri."
)
add_numbered_item(2,
    "Penggunaan asisten AI publik seperti ChatGPT tidak efektif sebagai solusi karena "
    "sistem tersebut tidak memiliki akses terhadap data internal SMKN 2 Indramayu, "
    "sehingga rentan menghasilkan jawaban yang tidak relevan, tidak akurat, atau bahkan "
    "keliru terkait informasi spesifik sekolah."
)
add_numbered_item(3,
    "Pertanyaan umum dari siswa baru saat ini dijawab secara manual oleh staf Tata Usaha "
    "setiap harinya, yang menyebabkan keterlambatan respons serta tidak adanya skalabilitas "
    "dalam pengelolaan informasi sekolah."
)

# 1.3 Tujuan
add_heading_left("1.3  Tujuan")

add_body(
    "Kegiatan proyek akhir ini membangun suatu sistem chatbot asisten virtual berbasis web "
    "untuk mengatasi permasalahan aksesibilitas informasi di lingkungan SMKN 2 Indramayu "
    "dengan memanfaatkan teknologi kecerdasan buatan. Adapun tujuan yang ingin dicapai "
    "secara spesifik adalah sebagai berikut."
)

add_numbered_item(1,
    "Menyediakan chatbot asisten virtual bernama DamayAI yang dapat menjawab pertanyaan "
    "seputar SMKN 2 Indramayu, mulai dari profil sekolah, data staf dan guru, hingga "
    "kegiatan sekolah, secara interaktif dan real-time tanpa perlu menghubungi staf "
    "secara manual."
)
add_numbered_item(2,
    "Membangun knowledge base kustom milik sekolah melalui proses scraping website dan "
    "penginputan data manual, sehingga setiap jawaban yang diberikan selalu berlandaskan "
    "informasi resmi dan spesifik SMKN 2 Indramayu."
)
add_numbered_item(3,
    "Mengimplementasikan mekanisme Tiered Retrieval-Augmented Generation yang "
    "memprioritaskan data terverifikasi milik sekolah sebelum menarik informasi dari "
    "sumber lain, guna meminimalkan halusinasi AI dan meningkatkan kepercayaan pengguna "
    "terhadap jawaban yang diberikan."
)

# 1.4 Manfaat
add_heading_left("1.4  Manfaat")

add_body(
    "Manfaat yang diharapkan dari pengembangan sistem DamayAI Assistant adalah sebagai "
    "berikut."
)

add_numbered_item(1,
    "Bagi siswa baru SMKN 2 Indramayu, DamayAI memberikan kemudahan dalam mengakses "
    "informasi seputar sekolah secara mandiri, cepat, dan kapan saja tanpa harus menunggu "
    "respons manual dari staf. Hal ini diharapkan dapat mempersingkat masa adaptasi siswa "
    "baru terhadap lingkungan sekolah."
)
add_numbered_item(2,
    "Bagi staf Tata Usaha dan guru SMKN 2 Indramayu, sistem ini membantu mengurangi beban "
    "kerja berulang akibat menjawab pertanyaan yang serupa setiap harinya, sehingga staf "
    "dapat lebih berfokus pada tugas-tugas yang membutuhkan penanganan langsung dan bersifat "
    "lebih kompleks."
)
add_numbered_item(3,
    "Bagi institusi SMKN 2 Indramayu secara keseluruhan, DamayAI dapat menjadi sarana "
    "digitalisasi layanan informasi sekolah yang modern, terpusat, dan berbasis data "
    "terverifikasi, sekaligus menjadi bukti penerapan teknologi kecerdasan buatan di "
    "lingkungan sekolah menengah kejuruan di Indonesia."
)
add_numbered_item(4,
    "Bagi pengembang dan komunitas akademik, proyek ini memberikan kontribusi berupa "
    "implementasi nyata pendekatan Tiered RAG untuk chatbot informasi sekolah menengah "
    "kejuruan di Indonesia, yang dapat dijadikan referensi atau dasar pengembangan sistem "
    "serupa di institusi pendidikan lainnya."
)

doc.add_page_break()

# Save checkpoint
doc.save(os.path.join(os.path.dirname(__file__), "DamayAI_Dokumen_PPA_Revised_Formatted.docx"))
print("Part 1 complete: Cover + BAB I saved.")
